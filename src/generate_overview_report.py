"""Generate the comprehensive Vietnamese thesis-overview Word document.

This document is the entry-point reading for anyone — advisor, examiner,
collaborator — who has never seen the project before. It narrates the
project end-to-end from medical motivation → dataset → methodology →
everything done in Step 0 through Step 4, with embedded figures and tables.

Target length: 40-60 pages. Audience: a smart reader with no background in
diabetes, CGM, or deep learning, who needs to understand WHAT this thesis
is solving and HOW.

Run from project root:
    python src/generate_overview_report.py
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_ROOT / "reports" / "bao_cao_tong_quan.docx"

FIG_DIR = PROJECT_ROOT / "outputs" / "figures"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x4E, 0x7E)


def h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x3D, 0x5F, 0x91)


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)


def info_box(doc, text, color="E8F1FB"):
    """Single-cell shaded box for definitions / sidenotes."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_shading(cell, color)
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    doc.add_paragraph()


def add_two_col_table(doc, rows, headers=None, widths=(4.5, 11.5)):
    n_rows = len(rows) + (1 if headers else 0)
    t = doc.add_table(rows=n_rows, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for w_i, w in enumerate(widths):
        for r in t.rows:
            r.cells[w_i].width = Cm(w)
    if headers:
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(c, "1F3A5F")
        offset = 1
    else:
        offset = 0
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + offset].cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True
    doc.add_paragraph()


def add_multi_col_table(doc, rows, headers, widths_cm):
    n_rows = 1 + len(rows)
    t = doc.add_table(rows=n_rows, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for col, w in enumerate(widths_cm):
        for r in t.rows:
            r.cells[col].width = Cm(w)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "1F3A5F")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc, fig_filename, caption, width_cm=15.5):
    path = FIG_DIR / fig_filename
    if not path.exists():
        para(doc, f"[Hình không tìm thấy: {fig_filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    doc.add_paragraph()


# ===========================================================================
# Document content
# ===========================================================================


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    # =========================================================
    # TITLE PAGE
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO TỔNG QUAN LUẬN VĂN")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A Multimodal Deep Learning Approach for\n"
        "Short-Term Blood Glucose Forecasting in Type 1 Diabetes"
    )
    r.italic = True
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tổng kết tiến độ từ Step 0 → Step 4")
    r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Sinh viên thực hiện: Sơn\n"
        "Đề tài: Luận văn tốt nghiệp đại học\n"
        "Dataset: HUPA-UCM (Hidalgo et al., Data in Brief 2024)\n"
        "Ngày báo cáo: 18/05/2026"
    )
    r.font.size = Pt(11)

    doc.add_paragraph()
    info_box(
        doc,
        "Mục đích tài liệu này: cung cấp một cái nhìn tổng quan đầy đủ để người đọc lần đầu — "
        "giảng viên hướng dẫn, người phản biện, đồng nghiệp — có thể hiểu được toàn bộ đề tài đang giải quyết "
        "bài toán gì, dùng dữ liệu nào, các kỹ thuật đã áp dụng và lý do đằng sau từng quyết định. "
        "Tài liệu được viết theo lối kể chuyện: bắt đầu từ bối cảnh lâm sàng, rồi đến bài toán toán học, "
        "rồi đến dữ liệu, rồi đến từng bước đã thực hiện. Mọi thuật ngữ chuyên môn đều được giải thích lần đầu khi xuất hiện.",
    )

    doc.add_page_break()

    # =========================================================
    # MỤC LỤC
    # =========================================================
    h1(doc, "MỤC LỤC")
    toc = [
        ("PHẦN 1", "Bối cảnh khoa học và lâm sàng"),
        ("PHẦN 2", "Đề tài luận văn — bài toán cần giải"),
        ("PHẦN 3", "Bài toán Machine Learning — định nghĩa toán học"),
        ("PHẦN 4", "Dataset HUPA-UCM — bộ dữ liệu chính"),
        ("PHẦN 5", "Workflow nghiên cứu — 14 bước có thứ tự"),
        ("PHẦN 6", "Step 0 — Đọc tài liệu, khảo sát công trình trước đó"),
        ("PHẦN 7", "Step 1 — Hiểu dữ liệu"),
        ("PHẦN 8", "Step 2 — Khám phá dữ liệu (EDA)"),
        ("PHẦN 9", "Step 3 — Tiền xử lý dữ liệu"),
        ("PHẦN 10", "Step 4 — Feature engineering + lựa chọn feature"),
        ("PHẦN 11", "Kiến trúc output hiện tại — những gì đã có sẵn"),
        ("PHẦN 12", "Định hướng tiếp theo — Step 5 đến Step 13"),
        ("PHẦN 13", "Bảng tóm tắt nhanh"),
    ]
    for k, v in toc:
        bullet(doc, f"{k}. {v}")
    doc.add_page_break()

    # =========================================================
    # PHẦN 1: BỐI CẢNH
    # =========================================================
    h1(doc, "PHẦN 1. BỐI CẢNH KHOA HỌC VÀ LÂM SÀNG")

    h2(doc, "1.1 Đái tháo đường Type 1 là gì?")
    para(
        doc,
        "Đái tháo đường (diabetes) là một nhóm bệnh chuyển hoá đặc trưng bởi tình trạng đường huyết "
        "(blood glucose) cao kéo dài. Trong cơ thể khoẻ mạnh, tuyến tuỵ tiết ra hormone insulin để giúp tế bào "
        "hấp thụ glucose từ máu — giống như chìa khoá mở cửa cho glucose vào tế bào. Khi cơ chế này bị rối loạn, "
        "glucose tích tụ trong máu thay vì được sử dụng, gây ra hậu quả nghiêm trọng cho mạch máu, thần kinh, "
        "thận, mắt và tim trong dài hạn.",
    )
    para(
        doc,
        "Có hai loại đái tháo đường chính. Type 2 (chiếm ~90% bệnh nhân) thường xuất hiện ở người trưởng thành, "
        "do tế bào trở nên kháng insulin; có thể quản lý bằng chế độ ăn, vận động và thuốc uống. "
        "Type 1 (T1D, chiếm ~10%) là một bệnh tự miễn — hệ miễn dịch của bệnh nhân tự tấn công và phá huỷ "
        "các tế bào beta của tuỵ, nơi sản xuất insulin. Bệnh nhân T1D KHÔNG TỰ TIẾT ĐƯỢC insulin nữa và phải "
        "tiêm insulin nhân tạo suốt đời. Đa số bệnh nhân T1D được chẩn đoán trước 30 tuổi.",
    )
    info_box(
        doc,
        "Đề tài này tập trung vào BỆNH NHÂN T1D — nhóm phải kiểm soát glucose bằng cách tiêm insulin chủ động. "
        "Đây là nhóm khó kiểm soát nhất vì mỗi sai sót về liều insulin hoặc thời điểm tiêm có thể gây "
        "hạ đường huyết (hypoglycemia) nguy hiểm trong vài chục phút.",
    )

    h2(doc, "1.2 Tại sao kiểm soát đường huyết là chuyện sống còn?")
    para(
        doc,
        "Glucose trong máu cần được giữ trong một khoảng hẹp. Quá thấp (hypoglycemia, <70 mg/dL) hoặc quá cao "
        "(hyperglycemia, >180 mg/dL) đều có hậu quả ngay lập tức và dài hạn:",
    )
    bullet(
        doc,
        "Hạ đường huyết nghiêm trọng (<54 mg/dL): có thể gây co giật, mất ý thức, hôn mê, thậm chí tử vong "
        "trong vài giờ nếu không xử lý. Đây là tai biến CẤP NGUY HIỂM NHẤT của T1D.",
    )
    bullet(
        doc,
        "Tăng đường huyết kéo dài: gây tổn thương vi mạch ở mắt (retinopathy → mù), thận (nephropathy → suy thận), "
        "thần kinh (neuropathy → mất cảm giác chi). Cũng tăng nguy cơ đột quỵ và nhồi máu cơ tim.",
    )
    bullet(
        doc,
        "Time-in-Range (TIR): tỷ lệ thời gian glucose trong khoảng 70-180 mg/dL. "
        "Chuẩn quốc tế (Battelino et al. 2019): TIR > 70% là mục tiêu. Mỗi 1% TIR tăng giảm 8% nguy cơ biến chứng.",
    )

    h2(doc, "1.3 Continuous Glucose Monitor (CGM) — cuộc cách mạng giám sát glucose")
    para(
        doc,
        "Trước năm 2000, bệnh nhân T1D chỉ có cách đo glucose bằng máy chích đầu ngón tay, vài lần mỗi ngày — "
        "tức là chỉ có 4-6 'điểm dữ liệu' trong 24 giờ, không thể thấy được toàn bộ động học của đường huyết.",
    )
    para(
        doc,
        "CGM (Continuous Glucose Monitor) thay đổi điều đó: một sensor nhỏ gắn dưới da liên tục đo glucose dịch kẽ "
        "(interstitial glucose) mỗi 1-15 phút. Trong dataset của thesis này, thiết bị FreeStyle Libre 2 đo mỗi 15 phút, "
        "tạo ra ~96 điểm/ngày. Đây là dữ liệu đủ dày để khảo sát PATTERN và DỰ BÁO XU HƯỚNG.",
    )

    h2(doc, "1.4 Bài toán: làm sao DỰ BÁO glucose trước khi nó hạ?")
    para(
        doc,
        "Vấn đề thực tế: bệnh nhân T1D vẫn liên tục bị hypo, đặc biệt là khi ngủ. Một cảnh báo chỉ kích hoạt KHI "
        "glucose đã hạ là quá muộn. Câu hỏi nghiên cứu là:",
    )
    info_box(
        doc,
        "Liệu có thể dùng dữ liệu lịch sử 30-120 phút gần nhất (glucose, insulin, ăn, vận động, nhịp tim...) "
        "để dự báo glucose ở thời điểm 30, 60, 90 phút TỚI với sai số đủ thấp, đủ để cảnh báo bệnh nhân "
        "TRƯỚC KHI hypo xảy ra?",
    )
    para(
        doc,
        "Nếu trả lời được câu hỏi này, ta có thể: xây dựng hệ thống cảnh báo sớm, hỗ trợ artificial pancreas "
        "(tuỵ nhân tạo tự động điều chỉnh insulin), tích hợp vào app điện thoại của bệnh nhân. "
        "Đây là lý do hàng trăm nghiên cứu trong 5 năm qua tập trung vào BLOOD GLUCOSE FORECASTING bằng machine learning.",
    )
    doc.add_page_break()

    # =========================================================
    # PHẦN 2: ĐỀ TÀI
    # =========================================================
    h1(doc, "PHẦN 2. ĐỀ TÀI LUẬN VĂN — BÀI TOÁN CẦN GIẢI")

    h2(doc, "2.1 Tên đề tài chính thức")
    para(doc, "A Multimodal Deep Learning Approach for Short-Term Blood Glucose Forecasting in Type 1 Diabetes", bold=True, italic=True, size=12)
    para(doc, "(Tiếp cận học sâu đa modal cho dự báo đường huyết ngắn hạn ở bệnh nhân Đái tháo đường Type 1)")

    h2(doc, "2.2 Phân tích từng từ trong tên đề tài")
    add_two_col_table(
        doc,
        rows=[
            ("Multimodal", "Sử dụng nhiều loại tín hiệu khác nhau: glucose (CGM), insulin tiêm (basal + bolus), "
             "carbohydrate ăn vào, vận động (bước chân, calories), nhịp tim. Multi-modal nghĩa là dữ liệu đa kênh "
             "đa nguồn — đối lập với chỉ dùng glucose đơn lẻ."),
            ("Deep Learning", "Sử dụng neural network có nhiều layer (sâu) — cụ thể là kiến trúc lai (hybrid) "
             "kết hợp CNN (Convolutional) cho local pattern + GRU (Gated Recurrent Unit) cho long-range memory + "
             "Attention để gate thông tin theo bệnh nhân."),
            ("Short-Term Forecasting", "Dự báo trong tương lai gần: 30, 60, 90 phút. Đây là khoảng thời gian "
             "có ý nghĩa lâm sàng — đủ dài để bệnh nhân có thời gian ăn carb chống hypo, đủ ngắn để dự báo còn chính xác."),
            ("Type 1 Diabetes", "Đối tượng nghiên cứu cụ thể: bệnh nhân T1D, dùng insulin chủ động (CSII pump hoặc MDI injection)."),
        ],
        headers=["Thuật ngữ", "Ý nghĩa trong đề tài"],
        widths=(3.5, 12.5),
    )

    h2(doc, "2.3 Đóng góp khoa học dự kiến")
    para(doc, "Đề tài này dự kiến đóng góp 5 điểm so với các công trình trước:", bold=True)
    bullet(
        doc,
        "1. Multimodal nghiêm túc: sử dụng đầy đủ 7 modality của HUPA-UCM (glucose, calories, HR, steps, basal, bolus, carb) "
        "thay vì chỉ dùng glucose như nhiều công trình trước.",
    )
    bullet(
        doc,
        "2. Patient embedding: thiết kế nhánh static branch riêng để model học cơ địa từng bệnh nhân (HbA1c, age, BMI, "
        "behavioural fingerprint từ train portion) — không phải one-model-fits-all.",
    )
    bullet(
        doc,
        "3. Cross-attention / gated fusion: cơ chế fusion linh hoạt giữa nhánh temporal và nhánh static, "
        "khác với cách concat đơn giản trong nhiều baseline.",
    )
    bullet(
        doc,
        "4. Xử lý vấn đề missing modality và sensor caps một cách methodology rõ ràng (modality availability flags, "
        "censoring flags) — không zero-fill silently như các paper trước.",
    )
    bullet(
        doc,
        "5. XAI và uncertainty: dùng SHAP + attention + integrated gradients để giải thích model, kèm Monte Carlo Dropout "
        "hoặc Conformal Prediction để cho ra prediction interval thay vì point estimate.",
    )

    h2(doc, "2.4 Output cuối cùng của đề tài (sau Step 13)")
    bullet(doc, "Model hybrid CNN-GRU-Attention đạt RMSE ≈ 15-20 mg/dL ở horizon 30 phút, ≈ 25-35 mg/dL ở 90 phút.")
    bullet(doc, "Báo cáo so sánh với 5-7 baselines: persistence, Ridge, Random Forest, XGBoost, LSTM, GRU đơn.")
    bullet(doc, "Hệ thống XAI giải thích từng dự báo: feature nào đóng góp bao nhiêu, attention vào thời điểm nào.")
    bullet(doc, "Prediction interval (uncertainty band) ở 80% và 95% confidence.")
    bullet(doc, "Demo app web (Streamlit/Gradio) — bệnh nhân upload CGM data, nhận dự báo + cảnh báo + giải thích.")
    bullet(doc, "Báo cáo luận văn ~80 trang theo định dạng IEEE / APA chuẩn.")

    h2(doc, "2.5 Đề tài KHÔNG phải gì (boundary)")
    bullet(doc, "Đây KHÔNG phải hệ thống y tế được cấp phép — chỉ là decision-support / research artefact.")
    bullet(doc, "KHÔNG kê liều insulin, KHÔNG ra phán quyết lâm sàng. Mọi output đều có disclaimer.")
    bullet(doc, "KHÔNG tự thu thập dữ liệu raw — sử dụng HUPA-UCM public dataset có sẵn.")
    bullet(doc, "KHÔNG so sánh với artificial pancreas đầy đủ — chỉ là một thành phần dự báo (prediction component).")

    doc.add_page_break()

    # =========================================================
    # PHẦN 3: BÀI TOÁN ML
    # =========================================================
    h1(doc, "PHẦN 3. BÀI TOÁN MACHINE LEARNING — ĐỊNH NGHĨA TOÁN HỌC")

    h2(doc, "3.1 Phát biểu bài toán")
    info_box(
        doc,
        "Cho lịch sử của bệnh nhân trong 120 phút gần nhất (24 timesteps × nhiều modality) và metadata "
        "tĩnh của bệnh nhân (HbA1c, BMI, treatment...), dự báo giá trị glucose tại 3 horizon tương lai: "
        "30 phút (t+6), 60 phút (t+12), và 90 phút (t+18). Đây là bài toán SUPERVISED MULTI-TASK REGRESSION.",
    )

    h2(doc, "3.2 Biểu diễn toán học")
    para(doc, "Input cho mỗi sample (một sequence) gồm 2 thành phần:", bold=True)
    bullet(doc, "X_dynamic ∈ R^(L × F) với L=24 (lookback), F=34 (số dynamic features). Đây là 'phim' về trạng thái patient.")
    bullet(doc, "X_static ∈ R^S với S=25 (số static features). Đây là 'CMND' cố định của patient.")
    para(doc, "Output cho mỗi sample:", bold=True)
    bullet(doc, "y ∈ R^3 = [glucose(t+6), glucose(t+12), glucose(t+18)] — 3 horizon được dự báo cùng lúc (multi-task).")

    h2(doc, "3.3 Hàm mất mát (loss function)")
    para(
        doc,
        "Hàm loss chính sẽ là Mean Squared Error (MSE) tổng trên 3 horizon: L = MSE(y_pred[0], y_true[0]) + "
        "MSE(y_pred[1], y_true[1]) + MSE(y_pred[2], y_true[2]). Trong Step 6, sẽ thêm hypo-penalty: nhân thêm "
        "trọng số cho lỗi ở zone hypoglycemia (<70 mg/dL) vì đây là zone clinically critical nhất.",
    )

    h2(doc, "3.4 Vì sao chọn 3 horizon 30/60/90 phút?")
    add_two_col_table(
        doc,
        rows=[
            ("30 phút (6 bin)", "Horizon ngắn nhất có ý nghĩa lâm sàng. Đủ để bệnh nhân kịp ăn carb chống hypo. "
             "Dễ dự báo (Spearman r với glucose hiện tại ≈ 0.85). Là 'baseline easy' để tất cả model đạt được."),
            ("60 phút (12 bin)", "Horizon trung — phổ biến nhất trong CGM forecasting literature. Vẫn dự báo được "
             "với sai số chấp nhận được (~25 mg/dL RMSE). Là horizon CHUẨN để so sánh model."),
            ("90 phút (18 bin)", "Horizon dài — challenging nhất, đo khả năng model học long-range dependency. "
             "Sai số tăng đáng kể (~35 mg/dL RMSE). Là benchmark để evaluate model có dùng tốt modality vận động/insulin không."),
        ],
        headers=["Horizon", "Vì sao chọn"],
        widths=(3.0, 13.0),
    )

    h2(doc, "3.5 Vì sao dùng nhiều modality thay vì chỉ glucose?")
    para(
        doc,
        "Một câu hỏi tự nhiên: glucose tự nó có autocorrelation rất cao — vậy đưa thêm insulin, carb, vận động "
        "có thực sự giúp không? Câu trả lời từ EDA (xem Phần 8) là CÓ, nhưng theo cách không trực tiếp:",
    )
    bullet(
        doc,
        "Glucose-only model có thể dự báo 'baseline drift' (trend) rất tốt, nhưng KHÔNG dự báo được 'shocks' "
        "như bữa ăn lớn (làm glucose tăng đột ngột) hoặc bolus correction (làm glucose hạ đột ngột).",
    )
    bullet(
        doc,
        "EDA peri-event analysis cho thấy: bolus correction đơn lẻ (không kèm carb) làm glucose giảm -30 mg/dL ở "
        "horizon 120 phút; carb đơn lẻ (không kèm bolus) làm glucose tăng +26 mg/dL. Đây là tín hiệu HUGE mà "
        "glucose-only model bỏ lỡ hoàn toàn.",
    )
    bullet(
        doc,
        "Vận động cao (>100 bước/5 phút) làm glucose giảm trung bình -4.8 mg/dL ở horizon 120 phút — tín hiệu "
        "yếu hơn nhưng đáng giữ.",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 4: DATASET
    # =========================================================
    h1(doc, "PHẦN 4. DATASET HUPA-UCM — BỘ DỮ LIỆU CHÍNH")

    h2(doc, "4.1 Nguồn gốc dataset")
    para(
        doc,
        "Tên đầy đủ: HUPA-UCM Diabetes Dataset. Được thu thập bởi Hospital Universitario Príncipe de Asturias "
        "(Alcalá de Henares, Tây Ban Nha) phối hợp Universidad Complutense de Madrid. Công bố trên Data in Brief "
        "(Elsevier) năm 2024 bởi Hidalgo, Alvarado, Botella et al. DOI: 10.17632/3hbcscwz44.1.",
    )
    para(
        doc,
        "Đây là một trong những public dataset T1D nghiêm túc nhất hiện có — có ethics approval "
        "(Protocol EC/11/2018), bệnh nhân ký consent, dữ liệu được làm sạch và publish dưới CC BY-NC-ND license.",
    )

    h2(doc, "4.2 Thông số cohort")
    add_multi_col_table(
        doc,
        rows=[
            ("Số bệnh nhân", "25", "Sau khi loại HUPA0008, HUPA0012, HUPA0013 (gaps quá lớn), còn 25."),
            ("Tỷ lệ giới tính", "52% nữ, 48% nam", "13 Female / 12 Male"),
            ("Tuổi (mean ± SD)", "39.2 ± 11.8", "Range 18.0 - 61.8"),
            ("HbA1c (mean ± SD)", "7.37 ± 0.82 %", "Range 6.0 - 9.7 (target lâm sàng <7%)"),
            ("Cân nặng", "69.1 ± 14.1 kg", "Range 51.0 - 104.8"),
            ("Chiều cao", "169.0 ± 10.4 cm", "Range 153 - 188"),
            ("Thời gian mắc bệnh", "17.8 ± 10.5 năm", "Range 0.8 - 39.5"),
            ("Treatment", "56% CSII / 44% MDI", "14 dùng pump CSII / 11 dùng MDI injection (sau khi P11 override)"),
            ("Tổng dòng dữ liệu", "309,392", "Trên 5-min grid"),
            ("Thời gian thu thập", "2018-06 đến 2022-05", "~4 năm, trải qua giai đoạn COVID lockdown"),
            ("Sampling cadence", "5 phút (sau preprocessing)", "Native: glucose 15 phút, Fitbit 1 phút"),
            ("Đơn vị glucose", "mg/dL", "Range 40 - 444 (sau khi censoring)"),
        ],
        headers=["Thông số", "Giá trị", "Ghi chú"],
        widths_cm=[3.5, 4.0, 8.5],
    )

    h2(doc, "4.3 Bảy modality được thu thập")
    add_multi_col_table(
        doc,
        rows=[
            ("glucose", "mg/dL", "FreeStyle Libre 2 CGM sensor (vùng cánh tay trên)", "Đo glucose dịch kẽ mỗi 15 phút."),
            ("calories", "kcal", "Fitbit Ionic smartwatch", "Calories đốt trong 5-min bin."),
            ("heart_rate", "bpm", "Fitbit Ionic optical sensor", "Nhịp tim trung bình trong bin."),
            ("steps", "số bước", "Fitbit Ionic accelerometer", "Tổng bước đi trong bin."),
            ("basal_rate", "U", "CSII pump (Medtronic/Roche) hoặc MDI app", "Insulin nền — continuous cho CSII, divided cho MDI."),
            ("bolus_volume_delivered", "U", "CSII pump hoặc MDI app", "Insulin nhanh tiêm thêm khi ăn hoặc correction."),
            ("carb_input", "servings (1 serving = 10g)", "Bệnh nhân tự log qua app", "Lượng carbohydrate ăn vào, ước lượng bởi bệnh nhân."),
        ],
        headers=["Modality", "Đơn vị", "Thiết bị thu thập", "Ý nghĩa lâm sàng"],
        widths_cm=[3.2, 2.5, 4.5, 6.0],
    )

    h2(doc, "4.4 Đặc điểm quan trọng cần biết về dataset")
    para(doc, "Pitfall 1 — Mất cân bằng thời lượng nghiêm trọng:", bold=True)
    para(
        doc,
        "HUPA0027P (574 ngày) chiếm 53.43% tổng dòng dữ liệu, HUPA0026P (141 ngày) chiếm 13.12%, "
        "HUPA0028P (90 ngày) chiếm 8.37%. Top 3 này gộp lại = 74.93% dataset. Nếu không xử lý, "
        "model sẽ bị 'dominate' bởi 3 bệnh nhân này — học cách glucose của họ thay vì học chung. "
        "Cách xử lý (xem Phần 9.8): cap N_train_cap=5000 sequences/patient trên TRAIN, val/test giữ nguyên.",
    )
    para(doc, "Pitfall 2 — Sensor cap censoring:", bold=True)
    para(
        doc,
        "FreeStyle Libre 2 báo 'LO' khi glucose ≤ 40 mg/dL, 'HI' khi > 400 — không phải đo lường thực. "
        "Trong dataset, 'LO' đã được convert thành số 40.0 (0.38% bin), 'HI' thành >400 (0.04% bin). "
        "Cần thêm cờ glucose_low_cap và glucose_high_extreme để model biết.",
    )
    para(doc, "Pitfall 3 — Missing modality patterns:", bold=True)
    para(
        doc,
        "5 bệnh nhân không có ít nhất 1 modality fully recorded: HUPA0011/0014/0015/0018/0020. "
        "4 bệnh nhân có basal coverage thấp 40-66%: HUPA0024/0026/0027/0028. "
        "Cần modality_available flags + basal_coverage_24h rolling.",
    )
    para(doc, "Pitfall 4 — Anomaly HUPA0011P:", bold=True)
    para(
        doc,
        "Metadata ghi HUPA0011 là CSII (dùng pump), nhưng dữ liệu KHÔNG có basal record nào. "
        "Pump CSII theo định nghĩa phải có basal liên tục. → Đây là data anomaly. "
        "Đã override treatment thành MDI trong static table để tránh model học rule sai 'CSII = no basal'.",
    )
    para(doc, "Pitfall 5 — COVID confound:", bold=True)
    para(
        doc,
        "HUPA0026/0027/0028 được thu thập từ 2020-05 trở đi, trùng giai đoạn lockdown Tây Ban Nha. "
        "Lifestyle, activity, meal patterns khác biệt rõ rệt so với 22 patient pre-COVID. "
        "Document làm limitation, không loại bỏ.",
    )

    add_figure(doc, "01_data_understanding_overview.png",
               "Hình 4.1: Tổng quan 25 bệnh nhân — thời lượng theo dõi, phân bố zone glucose, mất cân bằng dataset.")
    doc.add_page_break()

    # =========================================================
    # PHẦN 5: WORKFLOW
    # =========================================================
    h1(doc, "PHẦN 5. WORKFLOW NGHIÊN CỨU — 14 BƯỚC CÓ THỨ TỰ")

    para(
        doc,
        "Đề tài được tổ chức theo workflow 14 bước có thứ tự nghiêm ngặt (định nghĩa trong skills/SKILL.md). "
        "Mỗi step phải hoàn thành trước khi step sau bắt đầu, vì output của step trước là input cho step sau, "
        "và quan trọng hơn: mỗi quyết định methodology phải có evidence từ step trước.",
    )

    add_multi_col_table(
        doc,
        rows=[
            ("Step 0", "Literature review — đọc tài liệu, khảo sát công trình trước", "✅ Hoàn thành (có gaps)"),
            ("Step 1", "Đọc CLAUDE.md, SKILL.md, inspect data", "✅ Hoàn thành"),
            ("Step 2", "Meaningful EDA — khám phá dữ liệu", "✅ Hoàn thành (2 passes)"),
            ("Step 3", "Preprocessing — tiền xử lý", "✅ Hoàn thành"),
            ("Step 4", "Feature engineering + selection", "✅ Hoàn thành"),
            ("Step 5", "Baseline ladder — persistence, Ridge, RF/XGBoost, LSTM/GRU", "⏳ Chưa làm"),
            ("Step 6", "Hybrid model design + training", "⏳ Chưa làm"),
            ("Step 7", "Comparative analysis hybrid vs baselines", "⏳ Chưa làm"),
            ("Step 8", "Fine-tuning + ablation studies", "⏳ Chưa làm"),
            ("Step 9", "XAI integration (SHAP, attention, IG)", "⏳ Chưa làm"),
            ("Step 10", "Extended contributions — uncertainty, app, LLM", "⏳ Chưa làm"),
            ("Step 11", "Final report.md consolidation", "⏳ Chưa làm"),
            ("Step 12", "Colab cleanup + end-to-end test", "⏳ Chưa làm"),
            ("Step 13", "Application deployment (Streamlit/mobile)", "⏳ Chưa làm"),
        ],
        headers=["Step", "Mô tả", "Trạng thái"],
        widths_cm=[1.8, 11.5, 2.7],
    )

    info_box(
        doc,
        "Quy tắc vàng (SKILL.md Rule 1): KHÔNG SUY ĐOÁN trước khi inspect dữ liệu. Mọi quyết định "
        "methodology phải dựa trên dữ liệu thực hoặc tài liệu peer-reviewed. Đây là lý do thứ tự bước "
        "không thể đảo ngược: không thể chọn architecture model trước khi biết feature nào quan trọng, "
        "không thể design feature trước khi hiểu data.",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 6: STEP 0
    # =========================================================
    h1(doc, "PHẦN 6. STEP 0 — ĐỌC TÀI LIỆU, KHẢO SÁT CÔNG TRÌNH TRƯỚC ĐÓ")

    h2(doc, "6.1 Mục tiêu của step này")
    para(
        doc,
        "Trước khi bắt tay vào code, ta cần biết: (a) các công trình glucose forecasting trên HUPA-UCM trước đây "
        "đã dùng kỹ thuật gì, đạt kết quả gì; (b) các kỹ thuật chung trong CGM forecasting literature; "
        "(c) gap nào còn để thesis này lấp.",
    )

    h2(doc, "6.2 6 công trình HUPA-UCM trước đó (theo Hidalgo et al. 2024)")
    add_multi_col_table(
        doc,
        rows=[
            ("Tena 2021 (Sensors)", "Ensemble model — cutting-edge DNN", "Có thể dùng OhioT1DM thay vì HUPA — [verify]"),
            ("Alvarado 2023 (CILS)", "Wavelet + CNN cho hypoglycemia prediction", "Direct HUPA baseline cho hypo task"),
            ("Parra 2024 (IEEE JBHI)", "Structured Grammatical Evolution + Difference Eq", "Postprandial-specific"),
            ("Ingelse 2023 (Research Sq)", "Grammar-guided Genetic Programming", "Có thể 10-patient cohort khác HUPA-UCM"),
            ("Tena 2023 FPGA (IEEE JBHI)", "LSTM trên Xilinx FPGA cho deploy mobile, horizon 30 phút", "Direct baseline gần nhất cho thesis này"),
            ("Botella-Serrano 2023 (Frontiers)", "Sleep × glycemic control analysis", "Focus on sleep modality"),
        ],
        headers=["Paper", "Kỹ thuật chính", "Liên quan đề tài"],
        widths_cm=[4.5, 5.5, 6.0],
    )

    h2(doc, "6.3 Trạng thái Step 0 hiện tại")
    para(
        doc,
        "File reports/literature_review.md đã được draft (~27 KB) với coverage 3/6 paper HUPA. "
        "Còn 3 paper missing (Ingelse 2023, Tena 2023 FPGA, Botella-Serrano 2023), 4 citation cần verify, "
        "thiếu Gap Analysis section. Đây là known gap được track trong memory literature-review-gaps.md, "
        "sẽ được hoàn thiện trước khi viết §2 Background của report cuối.",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 7: STEP 1
    # =========================================================
    h1(doc, "PHẦN 7. STEP 1 — HIỂU DỮ LIỆU")

    h2(doc, "7.1 Quy trình inspecting")
    para(
        doc,
        "Step 1 không có 'kỹ thuật ML', chỉ là việc đọc dữ liệu một cách có hệ thống. "
        "Notebook 01_data_understanding.ipynb thực hiện:",
    )
    bullet(doc, "Load lần lượt 25 file Excel, kiểm tra schema 8 cột giống nhau.")
    bullet(doc, "Tính cohort summary: cho mỗi patient, đếm số rows, thời lượng, mean/std/min/max glucose, % thời gian ở 3 zones, censoring rate, modality coverage.")
    bullet(doc, "Verify timestamp: strict 5-min grid, monotonic, không duplicate.")
    bullet(doc, "Load patient_data_characteristic.xlsx, verify static metadata cho cả 25 patient.")
    bullet(doc, "Compute BMI từ weight/height.")
    bullet(doc, "Sinh figure tổng quan + bảng cohort_summary.csv.")

    h2(doc, "7.2 Output Step 1")
    bullet(doc, "data/interim/hupa_cohort_summary.csv — bảng 25 dòng × 22 cột (tham chiếu master cho mọi numeric claim sau).")
    bullet(doc, "outputs/figures/01_data_understanding_overview.png — figure tổng hợp 6 panel.")
    bullet(doc, "reports/report.md §3 'Dataset Description' — đã viết xong.")

    h2(doc, "7.3 Quyết định methodology từ Step 1")
    bullet(doc, "Glucose unit: mg/dL — confirmed.")
    bullet(doc, "Sampling: strict 5-min — không cần re-resample.")
    bullet(doc, "Cohort imbalance: top-3 patient = 74.93% rows → cần xử lý ở Step 3 (cap với adaptive stride).")
    bullet(doc, "Sensor caps: 40 và >400 cần flag, không drop.")
    bullet(doc, "BMI: tự tính, range 18.5-30.6.")

    doc.add_page_break()

    # =========================================================
    # PHẦN 8: STEP 2 EDA
    # =========================================================
    h1(doc, "PHẦN 8. STEP 2 — KHÁM PHÁ DỮ LIỆU (EXPLORATORY DATA ANALYSIS)")

    h2(doc, "8.1 Triết lý EDA của thesis")
    info_box(
        doc,
        "EDA không phải vẽ chart đẹp để báo cáo. EDA phải TRẢ LỜI CÂU HỎI cụ thể về methodology: "
        "Lookback length nên bao nhiêu? Modality nào đáng giữ? Có circadian effect không? "
        "Mỗi figure đều có một câu hỏi đứng sau và một implication cho Step 3.",
    )

    h2(doc, "8.2 9 sub-analysis trong EDA")

    h3(doc, "8.2.1 Phân bố glucose toàn cohort và per-patient")
    para(
        doc,
        "Câu hỏi: glucose có phân bố như thế nào? Có lệch hyperglycemia không?",
    )
    para(
        doc,
        "Findings: Mean 130 mg/dL (median), TIR cohort 60.7% (patient-averaged) vs 71.7% (row-weighted) — chênh "
        "11 điểm % do HUPA0027 dominance. Phân bố lệch phải (right-skewed) vì hyperglycemia frequent hơn hypo.",
    )
    para(doc, "Implication: TIR là metric chính, cần báo cáo cả 2 cách weighting song song.")
    add_figure(doc, "02_eda_glucose_distribution.png",
               "Hình 8.1: Phân bố glucose cohort-wide với 3 ngưỡng hypo/TIR/hyper.")

    h3(doc, "8.2.2 ACF/PACF — Autocorrelation Function")
    para(
        doc,
        "Câu hỏi: glucose hiện tại còn correlate với glucose trong quá khứ bao xa?",
    )
    para(
        doc,
        "ACF (Autocorrelation Function) đo correlation giữa series với chính nó shifted theo lag k. "
        "PACF (Partial ACF) trừ đi ảnh hưởng của các lag trung gian.",
    )
    para(
        doc,
        "Findings: ACF half-life trung bình 130 phút across 25 patients (range 75-260 phút). PACF tail off "
        "nhanh sau lag 4-6 (20-30 phút), suggesting AR(6) là baseline statistical model.",
    )
    para(doc, "Implication: Lookback 120 phút (24 bin) là hợp lý — bắt được phần lớn linear dependency.")
    add_figure(doc, "02_eda_acf.png", "Hình 8.2: Mean ACF của glucose trên cohort, với 95% CI band.")

    h3(doc, "8.2.3 Circadian pattern — chu kỳ theo giờ trong ngày")
    para(
        doc,
        "Câu hỏi: glucose có tăng/giảm theo giờ trong ngày không?",
    )
    para(
        doc,
        "Findings: Có chu kỳ rõ rệt — peak 07-10h sáng (~165 mg/dL, do 'dawn phenomenon' — cortisol kích hoạt "
        "gluconeogenesis vào sáng sớm), nadir 03-05h sáng (~135 mg/dL, do sleep deep + insulin sensitivity tăng).",
    )
    para(doc, "Implication: Time-of-day cyclical encoding (hour_sin/cos) là feature có giá trị.")
    add_figure(doc, "02_eda_circadian_profile.png", "Hình 8.3: Mean glucose theo giờ trong ngày, cohort-wide với IQR band.")

    h3(doc, "8.2.4 Peri-event analysis — phản ứng glucose sau bolus/carb/vận động")
    para(
        doc,
        "Câu hỏi: insulin có thực sự hạ glucose không? Carb có tăng glucose không? Vận động có ảnh hưởng không?",
    )
    para(
        doc,
        "Phương pháp: với mỗi event (bolus > 0, carb > 0, steps > 99), trích Δglucose tại t+6, t+12, t+18, t+24. "
        "So sánh với same-patient control sampling (rút random từ non-event timestamps cùng patient).",
    )
    para(doc, "Findings tổng quan:")
    bullet(doc, "Bolus tiêm (n=3,586): Δglucose -15.5 mg/dL @ 120 phút vs control (Δ-1.4). Effect rõ rệt.")
    bullet(doc, "Carb ăn (n=2,641): Δglucose +15.1 mg/dL @ 120 phút vs control. Effect rõ rệt.")
    bullet(doc, "Steps cao (n=29,870): Δglucose -4.8 mg/dL @ 120 phút vs control. Effect nhỏ hơn nhưng có.")
    add_figure(doc, "02_eda_peri_event.png", "Hình 8.4: Trajectory Δglucose sau bolus/carb/steps event vs control.")

    h3(doc, "8.2.5 Peri-event by subtype — tách bolus thành correction vs meal-bolus")
    para(
        doc,
        "Phát hiện thú vị Pass 2: cùng 1 bolus, nhưng effect khác biệt rõ rệt tuỳ context.",
    )
    bullet(doc, "Correction bolus (bolus đơn, không kèm carb): Δ -30 mg/dL @ 120 phút — gấp đôi pooled estimate.")
    bullet(doc, "Meal bolus (kèm carb): Δ ~-7 mg/dL — yếu hơn nhiều vì bị carb bù trừ.")
    bullet(doc, "Solo carb (carb đơn, không bolus): Δ +26 mg/dL @ 120 phút — nguy hiểm cho hyperglycemia.")
    para(
        doc,
        "Implication QUAN TRỌNG: bolus và carb phải được model như 2 modality RIÊNG BIỆT, "
        "không phải merge thành 'meal event'. Đây là cơ sở cho cross-attention/gated fusion architecture.",
    )
    add_figure(doc, "02_eda_peri_event_subtypes.png", "Hình 8.5: Tách correction vs meal-bolus vs solo-carb.")

    h3(doc, "8.2.6 Per-patient heterogeneity — bệnh nhân khác nhau như thế nào?")
    para(
        doc,
        "Findings: SD của peri-event response across 25 patients là 20-23 mg/dL — gần GẤP 2 LẦN magnitude trung bình. "
        "Tức là mỗi bệnh nhân phản ứng với cùng 1 bolus khác nhau ĐÁNG KỂ.",
    )
    para(
        doc,
        "Implication: One-model-fits-all sẽ kém — cần patient-conditioned modulation (cross-attention với static embedding). "
        "Đây là bằng chứng EMPIRICAL cho 2-branch architecture, không phải lựa chọn stylistic.",
    )
    add_figure(doc, "02_eda_per_patient_heterogeneity.png", "Hình 8.6: Cross-patient variance của peri-event response.")

    h3(doc, "8.2.7 Day-of-week effects")
    para(doc, "Findings: TIR weekend cao hơn weekday ~3-4 điểm %, do lifestyle điều độ hơn cuối tuần.")
    para(doc, "Implication: dayofweek_sin/cos là feature có giá trị (nhưng yếu hơn hour_sin/cos).")

    h3(doc, "8.2.8 Velocity by zone — tốc độ thay đổi glucose theo zone")
    para(
        doc,
        "Phân tích Δglucose/5min trong từng zone. Findings: ở zone hypo, velocity âm thường đi kèm acceleration âm "
        "(đang lao dốc), trong khi ở zone hyper thường velocity dương + acceleration giảm (đang đảo chiều).",
    )
    para(doc, "Implication: velocity và acceleration đều có information value độc lập.")

    h3(doc, "8.2.9 Retirement of Pearson screen")
    para(
        doc,
        "Trong EDA pass 1, đã thử lagged Pearson correlation giữa bolus/carb/steps và future glucose. "
        "Kết quả: |r| < 0.09 ở mọi lag, gợi ý các modality 'vô dụng'. "
        "Sau khi peri-event analysis cho thấy effect mạnh, ta hiểu Pearson r THẤP STRUCTURALLY: "
        "vì bolus/carb sparse 99%+ là zero, dot product mostly zero × non-zero = near-zero r.",
    )
    info_box(
        doc,
        "Đây là bài học methodology quan trọng: với sparse event streams (modality chỉ active ở <5% bin), "
        "Pearson r là metric SAI LẦM. Phải dùng peri-event analysis với same-patient control. "
        "Đã retire toàn bộ lagged Pearson CSV+figure khỏi outputs/ trong pass 2.",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 9: STEP 3 PREPROCESSING
    # =========================================================
    h1(doc, "PHẦN 9. STEP 3 — TIỀN XỬ LÝ DỮ LIỆU")

    h2(doc, "9.1 Triết lý 2 lớp")
    para(
        doc,
        "Pipeline tiền xử lý chia làm 2 lớp:",
        bold=True,
    )
    bullet(
        doc,
        "Lớp A (đã làm sẵn bởi tác giả dataset Hidalgo et al.): chuyển 7 modality thô từ Libre 2 + Fitbit + Pump "
        "về cùng grid 5 phút. Đây là pipeline glUCModel mà tác giả implement. Thesis này KẾ THỪA outputs, "
        "không claim đóng góp về imputation.",
    )
    bullet(
        doc,
        "Lớp B (đóng góp methodology của thesis): cờ censoring, cờ availability, sửa anomaly P11, time features, "
        "glucose derivatives, rolling event aggregates, target shifting, chronological split, per-subject Z-score, "
        "sequence construction với train cap, static feature table.",
    )

    h2(doc, "9.2 Lớp A — 7 kênh đã được tác giả xử lý ra sao")
    add_multi_col_table(
        doc,
        rows=[
            ("glucose", "FreeStyle Libre 2", "~15 phút",
             "Round timestamp về mốc 5-min → subsample về 15-min grid → linear interpolation về 5-min (max 1h gap)."),
            ("bolus_volume_delivered", "Pump/MDI app", "Event-based",
             "Sum trong bin 5 phút (conserve total dose). Empty bin → 0."),
            ("basal_rate", "Pump/MDI app", "Continuous (CSII) / 1 injection/ngày (MDI)",
             "CSII: sum rate trong bin. MDI: chia liều long-acting cho 288 (24h × 12) → uniform per-bin. Gaps → 0."),
            ("carb_input", "Bệnh nhân tự log", "Event-based",
             "Sum trong bin → chia 10 (gram → servings). Gaps → 0."),
            ("heart_rate", "Fitbit", "Irregular 1-5 phút",
             "Round về 5-min + linear interpolation cho gap (HR liên tục, không zero-fill)."),
            ("calories", "Fitbit", "1 phút",
             "Sum 5 record 1-phút thành 1 bin 5-phút (conserve quantity). Empty → 0."),
            ("steps", "Fitbit", "1 phút", "Tương tự calories."),
        ],
        headers=["Kênh", "Thiết bị", "Cadence native", "Xử lý ra grid 5-min"],
        widths_cm=[3.2, 2.5, 2.8, 7.5],
    )

    h2(doc, "9.3 Lớp B — Đóng góp methodology của thesis")

    h3(doc, "9.3.1 Censoring flags (Pitfall 2)")
    para(
        doc,
        "glucose_low_cap = 1 nếu glucose ≤ 40 (LO sensor). glucose_high_extreme = 1 nếu > 400 (HI sensor). "
        "GIỮ giá trị raw (không drop/cap về NaN), nhưng cho model biết qua flag để có thể down-weight loss "
        "ở censored points hoặc evaluate separately.",
    )

    h3(doc, "9.3.2 Sửa anomaly HUPA0011P (Pitfall 4)")
    para(
        doc,
        "Static table: treatment override CSII → MDI cho HUPA0011P. Lý do: bệnh nhân này metadata ghi CSII "
        "nhưng dataset không có basal record nào → không thể là pump user thực sự. Override để model không "
        "học rule sai 'CSII = không có basal signal'.",
    )

    h3(doc, "9.3.3 Modality availability flags (Pitfall 3)")
    para(doc, "3 cờ per-patient:")
    bullet(doc, "basal_available = 1 nếu patient có bất kỳ basal_rate > 0 ở bất cứ đâu trong timeline.")
    bullet(doc, "bolus_available = 1 nếu có bất kỳ bolus > 0.")
    bullet(doc, "carb_available = 1 nếu có bất kỳ carb_input > 0.")
    para(doc, "Plus 1 cờ rolling (dynamic): basal_coverage_24h = % bin có basal > 0 trong 24h gần nhất. "
              "Phân biệt 3 chế độ: CSII liên tục (~1.0), MDI sparse (~0.05), partial coverage (0.4-0.66).")

    h3(doc, "9.3.4 Time features cyclical")
    para(doc, "hour_sin = sin(2π × hour / 24), hour_cos = cos(2π × hour / 24). Tương tự dayofweek_sin/cos. "
              "Encoding chu kỳ thay vì integer raw → không có discontinuity tại midnight hoặc Sunday-Monday boundary.")

    h3(doc, "9.3.5 Glucose-derived features")
    bullet(doc, "glucose_30m_mean, glucose_60m_mean, glucose_120m_mean — rolling mean multi-scale.")
    bullet(doc, "glucose_velocity = Δglucose / 5 (mg/dL/min).")
    bullet(doc, "glucose_acceleration = Δvelocity / 5 (second derivative).")
    bullet(doc, "Step 4 thêm: glucose_60m_std (volatility).")

    h3(doc, "9.3.6 Rolling event aggregates")
    bullet(doc, "bolus_{30,60,180}m_sum — span theo dược động học rapid-acting insulin.")
    bullet(doc, "carb_{60,180}m_sum — span theo tốc độ hấp thu meal (Hovorka 2004).")
    bullet(doc, "steps_{30,150}m_sum — short walking bouts + post-exercise window.")
    bullet(doc, "calories_30m_sum, heart_rate_30m_mean.")
    bullet(doc, "Step 4 thêm: insulin_on_board (IOB) và carbs_on_board (COB) — exponential decay thay vì rectangular.")

    h3(doc, "9.3.7 Target construction")
    para(doc, "target_30m = glucose.shift(-6), target_60m = glucose.shift(-12), target_90m = glucose.shift(-18). "
              "Shift PER PATIENT để không bao giờ leak target sang patient khác trong concatenated table.")

    h3(doc, "9.3.8 Chronological split với boundary buffer")
    para(
        doc,
        "Mỗi patient: 70% timeline đầu → train, 15% kế → val, 15% cuối → test. KHÔNG random split (vi phạm "
        "rule chống leakage thời gian).",
    )
    para(
        doc,
        "Boundary buffer: 18 row (= max horizon) bị drop tại mỗi ranh giới train/val và val/test, để target "
        "của row train cuối cùng không leak vào val. Tổng cost: 900 row (0.29%) dropped.",
    )

    h3(doc, "9.3.9 Per-subject Z-score (fit train only)")
    para(
        doc,
        "Continuous features (glucose, HR, calories, rolling means, velocity, accel) → Z-score per-subject "
        "(mean/std tính per-patient từ TRAIN portion only, áp dụng cho val/test). Vì inter-patient variance "
        "rất lớn (glucose mean 113-201 mg/dL, std 35-85). Sparse event features (bolus/carb/steps + rolling sums) "
        "→ log1p + global Z-score (cùng scale across patients).",
    )

    h3(doc, "9.3.10 Sequence construction với adaptive-stride cap")
    para(doc, "Sliding window: lookback 24 bước (120 phút), horizons {6, 12, 18}. Mỗi anchor index t sinh 1 sequence "
              "[t-23 : t] với targets glucose[t+6/+12/+18]. Window bị reject nếu chạm buffer hoặc có NaN target.")
    para(doc, "Train cap (xem [[long-patient-strategy]] memory): per-patient, nếu > 5000 train anchor → "
              "adaptive stride = floor(n_train/5000), lấy uniformly dọc timeline. Patient ngắn giữ stride=1.")
    para(doc, "Val/test: stride=1, giữ NGUYÊN tất cả anchor.")

    h3(doc, "9.3.11 Static feature table")
    para(doc, "25 cột: 6 clinical numeric + 12 derived (train-only) + 3 modality binary + 4 one-hot. "
              "Tất cả Z-scored across 25 patient (1 row/patient → không có leakage thời gian).")

    add_figure(doc, "03_preprocessing_sequences_per_patient.png",
               "Hình 9.1: Số sequence per-patient sau split + cap. 3 long patients (26/27/28) cap tại 5000 train. "
               "Val/test theo tỷ lệ timeline (long patients dominate val/test volume nhưng không train).")

    h2(doc, "9.4 Output Step 3")
    add_two_col_table(
        doc,
        rows=[
            ("hupa_5min_timestep.parquet", "309,392 × 37 — bảng feature đầy đủ per-timestep với split labels."),
            ("hupa_5min_sequences.npz", "Tensor bundle: X_dynamic (159,172 × 24 × 31), X_static (× 25), y (× 3) — sẵn để feed model."),
            ("hupa_static_features.csv", "25 × 26 — patient-level static features Z-scored."),
            ("scalers.json", "Per-subject mean/std + global log1p params, fit trên train only."),
            ("hupa_split_boundaries.csv + summary.csv", "Bookkeeping: row indices của boundaries, sequence counts per-patient per-split."),
        ],
        headers=["File", "Mô tả"],
        widths=(5.5, 10.5),
    )
    info_box(
        doc,
        "Sau Step 3: 309,392 row → 159,172 sequence được phân hoạch train=68,395 / val=45,382 / test=45,395, "
        "mỗi sequence có shape input (24, 31) tensor + 25 static + 3 horizon target. "
        "Pipeline deterministic ở SEED=42, end-to-end ~50s local.",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 10: STEP 4
    # =========================================================
    h1(doc, "PHẦN 10. STEP 4 — FEATURE ENGINEERING + LỰA CHỌN FEATURE")
    para(
        doc,
        "Phần này được tổ chức theo logic top-down: (10.1) Vì sao cần feature engineering; (10.2) Phân loại 10 nhóm chức năng; "
        "(10.3) Từng nhóm: feature, công thức tính, ý nghĩa cho model; (10.4) Scaling và phòng leakage; "
        "(10.5) Selection analysis chính thức + 2 vòng pruning.",
    )

    h2(doc, "10.1 Vì sao cần feature engineering?")
    para(doc, "Ba thuộc tính của dataset HUPA-UCM khiến raw input không phù hợp để model trực tiếp tiêu thụ:")
    para(doc, "1. Cấu trúc thời gian đa-scale (multi-scale temporal):", bold=True)
    para(
        doc,
        "Glucose dynamics evolve song song trên nhiều timescale: bữa ăn raise glucose 30-120 phút (Hovorka 2004), "
        "rapid-acting insulin act 60-180 phút onset 15 phút (Mathieu 2017), exercise giảm glucose 1-2h sau khi xong "
        "(Riddell 2017), dawn phenomenon shift basal mức 24h. EDA §4.4: glucose ở 120 phút vẫn correlate 0.50 với hiện tại. "
        "Single representation tại bin hiện tại mất context multi-scale → engineered rolling means 30/60/120 phút + "
        "60-phút std làm các scale explicit.",
    )
    para(doc, "2. Sparse event streams encoding pharmacology non-linear:", bold=True)
    para(
        doc,
        "Bolus và carb là event-based: hầu hết bin = 0, vài bin non-zero encode event có effect phụ thuộc thời gian "
        "đã trôi qua + pharmacokinetics. Linear model không thể encode exponential decay kernel từ raw lookback; "
        "tree model không thể encode continuous weighting. Engineered pharmacokinetic decay features (IOB, COB) "
        "inject physiological prior trực tiếp vào input.",
    )
    para(doc, "3. Patient heterogeneity + missing-modality ambiguity:", bold=True)
    para(
        doc,
        "Per-patient glucose mean range 113-201 mg/dL, std 35-85 mg/dL (EDA §4.7). 5/25 bệnh nhân thiếu ≥1 modality "
        "(Pitfall #6). Không có engineered subject_mean_glucose / subject_hypo_pct / etc., model không phân biệt được "
        "'patient này thường cao' vs 'glucose hiện tại cao'. Không có availability flags, model không phân biệt "
        "'không có event ở bin này' vs 'patient này không bao giờ ghi modality này'.",
    )
    info_box(
        doc,
        "Engineering features encode 3 thuộc tính trên đạt 2 mục tiêu: (a) giảm inductive load của temporal model "
        "→ architecture nhỏ hơn fit cùng data; (b) cho tabular baselines (Ridge/RF/XGBoost) cơ hội fair trong "
        "comparative evaluation §7.",
    )

    h2(doc, "10.2 Phân loại 10 nhóm chức năng")
    para(
        doc,
        "33 features chia thành 17 dynamic + 16 static, mỗi loại tiếp tục chia subgroups theo CHỨC NĂNG "
        "(không phải theo nguồn raw measurement):",
    )
    add_multi_col_table(
        doc,
        rows=[
            ("A. Glucose representations", "Dynamic, 6", "glucose, glucose_30/60/120m_mean, glucose_60m_std, glucose_velocity",
             "Multi-scale view của target signal + trend + variability."),
            ("B. Physiological raw", "Dynamic, 2", "heart_rate, basal_rate",
             "Continuous device measurements, temporal trunk localise timing."),
            ("C. Event rolling aggregates", "Dynamic, 3", "bolus_60m_sum, steps_150m_sum, heart_rate_30m_mean",
             "Cumulative magnitudes trong span sinh lý — tree baselines cần explicit."),
            ("D. Pharmacokinetic decay", "Dynamic, 2", "insulin_on_board, carbs_on_board",
             "Recency-weighted cumulative active-drug, inject pharmacological priors."),
            ("E. Cyclical time", "Dynamic, 2", "hour_sin, hour_cos",
             "Smooth representation time-of-day cho circadian patterns."),
            ("F. Operational flags", "Dynamic, 2", "glucose_low_cap, basal_coverage_24h",
             "Báo cho model khi measurement bị censored hoặc modality partial coverage."),
            ("I. Clinical metadata", "Static, 4", "hba1c_pct, age_years, dx_time_years, bmi",
             "Patient-level clinical context entered tại enrolment."),
            ("II. Behavioral fingerprint", "Static, 7", "subject_mean/std/hypo/hyper_pct_glucose, bolus_events_per_day, steps_active_pct, mean_heart_rate",
             "Per-patient summary statistics tính từ TRAIN portion only — 'patient này thường thế nào'."),
            ("III. Modality availability", "Static, 3", "basal_available, bolus_available, carb_available",
             "Báo cho model patient này RECORD modality nào."),
            ("IV. Demographic one-hot", "Static, 2", "gender_Female, treatment_CSII",
             "Categorical patient context."),
        ],
        headers=["Nhóm", "Loại + số", "Members", "Chức năng trong model"],
        widths_cm=[3.5, 1.8, 5.5, 5.2],
    )

    h2(doc, "10.3 Chi tiết từng nhóm — feature, công thức, ý nghĩa")

    h3(doc, "10.3.1 Nhóm A — Glucose representations (6 dynamic)")
    para(
        doc,
        "Raw glucose tại bin hiện tại là feature dominant (Spearman r=+0.69 với target_60m, perm imp gấp 2 lần "
        "feature mạnh thứ 2). Ba rolling means (30/60/120 phút, công thức rolling(window=k).mean()) cung cấp "
        "multi-scale low-pass filtering. glucose_60m_std (rolling(12).std()) bắt được volatility — Battelino 2019 "
        "consensus đề xuất là metric variability chính. glucose_velocity = Δglucose/5 (mg/dL/phút) — same value "
        "với velocity -2 vs velocity 0 là trajectories khác hẳn (tương đương CGM trend arrows ↑↓).",
    )

    h3(doc, "10.3.2 Nhóm B — Physiological raw signals (2 dynamic)")
    para(
        doc,
        "heart_rate là Fitbit HR raw. Giữ raw (không chỉ rolling) vì HR bursts (stress, exercise onset) carry timing "
        "info mà rolling smooth ra. basal_rate là insulin nền tại bin hiện tại — CSII pumps tự log liên tục, MDI thì "
        "spread đều theo §5.1.3. Giữ raw vì pump users có thể temp basal trong thời gian ngắn.",
    )

    h3(doc, "10.3.3 Nhóm C — Event rolling aggregates (3 dynamic)")
    para(
        doc,
        "Mỗi rolling match action-time của biology underlying. bolus_60m_sum (rolling(12).sum()) tương ứng peak action "
        "của rapid-acting insulin — feature explicit cho tree baselines không reconstruct IOB kernel được. "
        "steps_150m_sum (rolling(30).sum()) span cả short walking + post-exercise insulin sensitivity 1-2h. "
        "heart_rate_30m_mean smooth HR raw để bắt baseline tone (sleep state, stress kéo dài).",
    )

    h3(doc, "10.3.4 Nhóm D — Pharmacokinetic decay (2 dynamic, Step 4 contribution)")
    para(
        doc,
        "Đây là đóng góp methodology mới ở Step 4. Vấn đề với rolling sum: kernel hình chữ nhật — bolus tiêm 5 phút "
        "trước và 179 phút trước có cùng weight = 1. Nhưng dược động học: bolus mới còn rất active, bolus cũ đã decay.",
    )
    info_box(
        doc,
        "IOB: insulin_on_board[t] = α × IOB[t-1] + bolus[t], α = exp(-Δt/τ) = exp(-5/75) = 0.9355.\n"
        "Tương đương: IOB[t] = Σ_{k=0}^∞ exp(-k×5/75) × bolus[t-k]. Sau 5τ=375 phút residual <1%.\n"
        "τ_IOB = 75 phút match rapid-acting insulin (lispro/aspart/glulisine, Mathieu 2017).\n\n"
        "COB: carbs_on_board[t] tương tự với τ_COB = 60 phút (Hovorka 2004 first-order absorption).",
    )
    para(
        doc,
        "Lợi ích: tree-based baselines truy cập trực tiếp 'insulin còn active right now' — pattern tree không tự học. "
        "Cho NN: giảm inductive load. IOB + COB tạo cặp đối lập (insulin hạ vs carb tăng) → model có thể học net effect.",
    )

    h3(doc, "10.3.5 Nhóm E — Cyclical time encoding (2 dynamic)")
    para(
        doc,
        "hour_sin = sin(2π × h / 24), hour_cos = cos(2π × h / 24). Đưa giờ trong ngày thành point trên unit circle. "
        "EDA §4.5 documented dawn phenomenon (peak 07-10h, do cortisol kích hoạt gluconeogenesis sớm) và nadir 03-05h. "
        "Sin-cos pair tránh discontinuity tại 23h59 → 00h00 mà integer-hour có. Dayofweek_sin/cos đã drop vì weekend "
        "effect chỉ 3-4 pp TIR (yếu, §10.5).",
    )

    h3(doc, "10.3.6 Nhóm F — Operational và censoring flags (2 dynamic)")
    para(
        doc,
        "glucose_low_cap = 1 nếu glucose ≤ 40 (FreeStyle Libre 2 LO representation). Cho loss có Tobit-style mask. "
        "basal_coverage_24h = rolling(288).mean() của (basal_rate > 0) — phân biệt 3 chế độ: CSII liên tục (~1.0), "
        "MDI sparse (~0.05), partial coverage 4 patient (0.4-0.66, §3.6.3). Khác static availability flags ở chỗ "
        "time-varying — bắt được sự kiện tháo pump tạm thời.",
    )

    h3(doc, "10.3.7 Nhóm I — Clinical metadata (4 static)")
    para(
        doc,
        "Bốn cột nhập một lần tại enrolment, không thay đổi. hba1c_pct: gold-standard 3-month average glucose biomarker, "
        "range 6.0-9.7% (target lâm sàng <7%). age_years (18-61.8): shift insulin sensitivity. dx_time_years (0.8-39.5): "
        "phân biệt honeymoon phase vs long-standing. bmi = weight/(height/100)² — summarise body-composition driven "
        "insulin resistance (weight + height đã drop vì redundant).",
    )

    h3(doc, "10.3.8 Nhóm II — Behavioral fingerprint từ TRAIN ONLY (7 static)")
    para(
        doc,
        "Đây là methodology contribution của thesis. Mỗi patient được tính 7 thống kê CHỈ TỪ TRAIN portion của họ, "
        "không bao giờ touched val/test — đảm bảo no leakage. Broadcast vào mọi sequence của patient đó:",
    )
    bullet(doc, "subject_mean_glucose — typical glucose level (range 113-201 mg/dL across cohort)")
    bullet(doc, "subject_std_glucose — typical variability (range 35-85)")
    bullet(doc, "subject_hypo_pct (TBR — Battelino target <4%)")
    bullet(doc, "subject_hyper_pct (TAR — target <25%)")
    bullet(doc, "bolus_events_per_day — pump users typical 4-7, MDI 3-4")
    bullet(doc, "steps_active_pct — % bin > 99 steps (intensity proxy)")
    bullet(doc, "mean_heart_rate — baseline cardiovascular tone")
    para(
        doc,
        "5/7 ranking top 10 composite — bằng chứng EMPIRICAL (không phải stylistic) cho việc dùng patient-embedding "
        "branch riêng.",
    )

    h3(doc, "10.3.9 Nhóm III — Modality availability (3 static)")
    para(
        doc,
        "basal_available, bolus_available, carb_available — binary per-patient (1 nếu patient có RECORD bất kỳ modality "
        "đó trong toàn timeline). Giải quyết structural ambiguity của 'zero': '0 = không có event ở bin' vs '0 = patient "
        "này không bao giờ record'. Không có flag, model học rằng 5 patient missing carb đơn giản không bao giờ ăn carb "
        "(biological nonsense).",
    )

    h3(doc, "10.3.10 Nhóm IV — Demographic categorical (2 static)")
    para(
        doc,
        "gender_Female và treatment_CSII là halves đã giữ của 2 one-hot pairs. Complementary halves (gender_Male, "
        "treatment_MDI) drop vì derivable = 1 - x. Gender shift insulin sensitivity qua menstrual cycle (Pickup 2014). "
        "Treatment_CSII là **fundamental clinical distinction** giữa pump (continuous basal, nhiều bolus) vs injection "
        "(sparse basal, larger event).",
    )

    h2(doc, "10.4 Scaling và phòng leakage")
    para(doc, "Mọi engineered feature cần scaling, và mọi scaling là nguy cơ train-test leakage nếu fit ngoài train.")
    add_multi_col_table(
        doc,
        rows=[
            ("Per-subject Z-score", "glucose, heart_rate, 3 glucose rolling means, glucose_60m_std, glucose_velocity",
             "Fit (mean, std) per-patient TỪ TRAIN ROWS ONLY. Apply cho val/test. Justify bởi inter-patient variance lớn "
             "(EDA §4.7: glucose mean 113-201 mg/dL across 25 patient)."),
            ("log1p + Global Z-score", "basal_rate, bolus_60m_sum, steps_150m_sum, heart_rate_30m_mean, IOB, COB",
             "log1p compress heavy tail của sparse events. Global Z-score fit trên train rows pooled. "
             "Justify: scale insulin/steps comparable across patients cùng therapy class."),
            ("Pass-through", "hour_sin, hour_cos, glucose_low_cap, basal_coverage_24h",
             "Đã bounded [-1,1] hoặc [0,1]. Không scale."),
            ("Static cohort Z-score", "16 numeric static features",
             "Z-score across 25-patient table. Leakage-safe vì mỗi patient 1 row, không có time-series ordering."),
        ],
        headers=["Scaling type", "Features áp dụng", "Lý do + cách phòng leakage"],
        widths_cm=[3.2, 6.5, 6.3],
    )
    para(doc, "Leakage prevention được enforce tại 4 điểm trong pipeline:", bold=True)
    bullet(doc, "(1) Chronological 70/15/15 split per-patient áp dụng TRƯỚC khi engineer feature.")
    bullet(doc, "(2) Rolling/pharmacokinetic compute WITHIN từng patient bằng groupby('participant_id').")
    bullet(doc, "(3) Group II behavioral statistics filter df[split=='train'] trước khi aggregate.")
    bullet(doc, "(4) Buffer 18 row tại mỗi split boundary chống target leak (chỉ 0.29% rows dropped).")

    h2(doc, "10.5 Feature selection analysis — 3 signals + 2-round pruning")
    para(
        doc,
        "Selection analysis chính thức (SKILL.md §4.5) evaluate mọi feature với 3 signal orthogonal trên train portion only:",
    )
    bullet(doc, "Spearman rank correlation — monotonic association robust outlier.")
    bullet(doc, "Mutual information (sklearn k-NN Kraskov 2004) — bắt non-linear non-monotonic.")
    bullet(doc, "Random Forest permutation importance (n_estimators=80, max_depth=12, R² train=0.83) — capture cross-feature interactions.")
    bullet(doc, "Composite rank = unweighted mean của 3 rank.")
    para(doc, "Vòng 1 — Information-theoretic redundancy (drop 21):", bold=True)
    add_multi_col_table(
        doc,
        rows=[
            ("A. Math redundant", "4", "weight/height (BMI cover), gender_Male/treatment_MDI (one-hot sum=1)"),
            ("B. IOB/COB redundant", "8", "raw bolus/carb (recoverable từ IOB/COB), 4 rolling sums bolus+carb"),
            ("C. Wearable redundant", "4", "raw calories, raw steps, mean_daily_steps, steps_30m_sum"),
            ("D. Statistical", "1", "subject_tir_pct = 100 - hypo - hyper"),
            ("E. Weak/duplicate", "5", "dayofweek_sin/cos, glucose_high_extreme (0.04%), 3 dynamic availability flags"),
        ],
        headers=["Category", "Số", "Features dropped"],
        widths_cm=[3.5, 1.5, 11.0],
    )
    para(doc, "Vòng 2 — Clinical-lens criterion (drop thêm 5):", bold=True)
    add_multi_col_table(
        doc,
        rows=[
            ("glucose_acceleration", "Không CGM device nào show 2nd derivative. Velocity đã capture actionable trend (trend arrows)."),
            ("calories_30m_sum", "Calories từ Fitbit HR+steps classifier; doctors không dùng calories cho glucose mgmt."),
            ("carb_events_per_day", "Đo tần suất LOG meals, không phải eating. Bệnh nhân không log → feature misleading."),
            ("data_duration_days", "Pure data artifact zero clinical meaning. Deploy bệnh nhân mới chỉ 7 ngày → tín hiệu sai."),
            ("basal_recording_pct", "Correlate ~0.9 với treatment_CSII. basal_available binary đã encode."),
        ],
        headers=["Feature dropped", "Lý do clinical"],
        widths_cm=[3.5, 12.5],
    )
    info_box(
        doc,
        "FINAL: 17 dynamic + 16 static = 33 features. X_dynamic (159172, 24, 17), X_static (159172, 16). "
        "Feature budget per sample = 24×17 + 16 = 424 (giảm 49.6% từ 841 gốc). "
        "Top 10 ranking: glucose dominate (rank 1, Spearman +0.69), 5/10 là static behavioral fingerprint "
        "(subject_mean_glucose rank 2, subject_hyper_pct rank 5, subject_std_glucose rank 6, subject_hypo_pct rank 9, "
        "steps_active_pct rank 10) → empirical justification cho patient-embedding architecture.",
    )
    add_figure(doc, "04_feature_selection_dynamic.png",
               "Hình 10.1: Top features dynamic theo 3 signal — glucose family dominate.")
    add_figure(doc, "04_feature_selection_static.png",
               "Hình 10.2: Static features ranking — behavioral fingerprint mạnh hơn metadata clinical.")

    doc.add_page_break()

    # =========================================================
    # PHẦN 11: ARTEFACTS HIỆN TẠI
    # =========================================================
    h1(doc, "PHẦN 11. KIẾN TRÚC OUTPUT HIỆN TẠI — NHỮNG GÌ ĐÃ CÓ SẴN")

    h2(doc, "11.1 Cấu trúc thư mục dự án")
    add_two_col_table(
        doc,
        rows=[
            ("data/data_hupa/", "Dataset gốc HUPA-UCM (25 .xlsx + patient_data_characteristic.xlsx + Article_data.md)"),
            ("data/processed/", "Output Step 3+4: parquet timestep, npz sequence bundle, csv static features"),
            ("data/interim/", "Cohort summary từ Step 1, EDA patient summary từ Step 2"),
            ("src/", "Module Python: config, data_loading, preprocessing, feature_selection, generate_*_report"),
            ("notebooks/", "Notebook Colab-compatible: 01_data_understanding, 02_eda, 03_preprocessing"),
            ("outputs/tables/", "15+ CSV với findings EDA + feature selection ranking"),
            ("outputs/figures/", "PNG plots (1 từ Step 1, 7 từ Step 2, 1 từ Step 3, 2 từ Step 4)"),
            ("outputs/models/", "scalers.json (sẽ thêm model checkpoints ở Step 5+)"),
            ("reports/", "report.md (báo cáo chính §1-§6 đã viết), literature_review.md, feature_catalogue.docx, "
                          "feature_selection_report.docx, bao_cao_tong_quan.docx (chính tài liệu này)"),
            ("skills/", "SKILL.md — protocol 14 bước"),
            ("CLAUDE.md", "Project-specific guidance cho Claude Code"),
        ],
        headers=["Đường dẫn", "Nội dung"],
        widths=(4.5, 11.5),
    )

    h2(doc, "11.2 Sequence bundle — input chính cho model")
    para(doc, "File data/processed/hupa_5min_sequences.npz chứa 10 array:")
    add_multi_col_table(
        doc,
        rows=[
            ("X_dynamic", "(159172, 24, 17) float32", "Input chuỗi 24 timesteps × 17 features (sau pruning 2 vòng từ 34). Đã Z-scored / log1p+scaled tuỳ feature."),
            ("X_static", "(159172, 16) float32", "Vector static 16-d per sequence (sau pruning từ 25). Broadcast từ static_table 16-row."),
            ("y", "(159172, 3) float32", "Target glucose tại horizon 30/60/90 phút. Ở scale gốc mg/dL (không scaled)."),
            ("participant_ids", "(159172,) object", "Tên patient (HUPA0001P...) cho mỗi sequence — phục vụ patient-level CV."),
            ("split", "(159172,) object", "'train'/'val'/'test' cho mỗi sequence."),
            ("anchor_time", "(159172,) int64", "Timestamp của bin cuối lookback — debug + leakage check."),
            ("feature_names_dynamic", "(17,) object", "Tên 17 dynamic features theo thứ tự axis -1 của X_dynamic."),
            ("feature_names_static", "(16,) object", "Tên 16 static features."),
            ("horizon_minutes", "(3,) int32", "[30, 60, 90] — reference."),
            ("lookback_steps", "scalar int32", "24 — reference."),
        ],
        headers=["Array", "Shape / dtype", "Ý nghĩa"],
        widths_cm=[3.5, 3.8, 8.7],
    )

    h2(doc, "11.3 Cách load và sử dụng (cho Step 5+)")
    info_box(
        doc,
        "import numpy as np\n"
        "data = np.load('data/processed/hupa_5min_sequences.npz', allow_pickle=True)\n"
        "X_dyn, X_stat, y = data['X_dynamic'], data['X_static'], data['y']\n"
        "split = data['split']\n"
        "X_train_dyn = X_dyn[split=='train']; y_train = y[split=='train']\n"
        "# Train model... evaluate trên val/test\n"
        "# Inverse-scale y_pred bằng scalers.json để tính RMSE ở mg/dL gốc",
    )

    doc.add_page_break()

    # =========================================================
    # PHẦN 12: ĐỊNH HƯỚNG TIẾP THEO
    # =========================================================
    h1(doc, "PHẦN 12. ĐỊNH HƯỚNG TIẾP THEO — STEP 5 ĐẾN STEP 13")

    h2(doc, "12.1 Step 5 — Baseline ladder (sẽ làm tiếp theo)")
    para(
        doc,
        "Mục tiêu: implement và đánh giá các baseline đơn giản TRƯỚC khi train hybrid model. "
        "Nếu hybrid không thắng baseline, không có contribution. Nếu thắng nhỏ, complexity không justified.",
    )
    add_multi_col_table(
        doc,
        rows=[
            ("Persistence", "y_pred(t+h) = glucose(t) — không model gì cả, chỉ giả định glucose không đổi.", "Baseline tuyệt đối thấp nhất."),
            ("Ridge regression", "Linear với L2, input = flatten(X_dyn) + X_stat = 841 features.", "Test xem linearity đi xa được không."),
            ("Random Forest", "n_estimators ~ 200-500, max_depth ~ 12-20.", "Capture cross-feature interaction. Strong baseline trong tabular."),
            ("XGBoost / LightGBM", "Gradient boosting trees, tune learning rate + depth.", "State-of-the-art tabular baseline."),
            ("LSTM single-branch", "Chỉ ăn X_dynamic, không dùng X_static. Hidden ~ 64-128.", "Test xem temporal model có cần static không."),
            ("GRU + static", "GRU ăn X_dyn + Dense ăn X_stat, concat fusion. Hidden 64.", "Closest baseline cho hybrid đề xuất."),
        ],
        headers=["Model", "Config", "Mục đích"],
        widths_cm=[3.2, 7.0, 5.8],
    )

    h2(doc, "12.2 Step 6 — Hybrid model đề xuất")
    para(doc, "Kiến trúc dự kiến:")
    bullet(doc, "Temporal branch: 1D-CNN (kernel size 3-5) → BiGRU 2 layers (hidden 64-128).")
    bullet(doc, "Static branch: MLP 2 layers Dense 64→32.")
    bullet(doc, "Fusion: Cross-attention (query=GRU output, key/value=static embedding) HOẶC gated fusion.")
    bullet(doc, "Output head: Dense → 3 outputs (multi-horizon).")
    bullet(doc, "Loss: MSE multi-task + hypo-penalty (weight 2× cho glucose < 70 mg/dL).")

    h2(doc, "12.3 Step 7 — So sánh hybrid vs baseline")
    para(doc, "Bảng metric per-horizon × per-model × per-zone. Confidence interval bằng bootstrap. "
              "Clarke Error Grid Analysis (EGA) — chuẩn lâm sàng cho glucose forecast.")

    h2(doc, "12.4 Step 8 — Ablation studies")
    para(doc, "Verify từng thiết kế đóng góp gì:")
    bullet(doc, "Drop static branch — hybrid vs hybrid-without-static.")
    bullet(doc, "Drop modality một-by-one — bolus, carb, activity (M0→M4 ladder).")
    bullet(doc, "Drop dynamic availability flags (redundant với static từ §6).")
    bullet(doc, "Drop raw event streams (chỉ giữ rolling versions).")
    bullet(doc, "Drop IOB/COB (verify pharm decay vs rectangular sum).")
    bullet(doc, "Lookback 24 vs 36 (120 vs 180 phút).")

    h2(doc, "12.5 Step 9 — XAI")
    bullet(doc, "SHAP: DeepExplainer trên hybrid model, GradientExplainer cho NN, TreeExplainer cho RF baseline.")
    bullet(doc, "Attention visualization: heatmap attention weights theo timestep + theo modality.")
    bullet(doc, "Integrated Gradients: feature attribution model-agnostic.")
    bullet(doc, "Permutation importance trên test set: cross-check SHAP.")

    h2(doc, "12.6 Step 10 — Extended contributions")
    bullet(doc, "Uncertainty quantification: MC Dropout, Conformal Prediction, Deep Ensembles.")
    bullet(doc, "App demo Streamlit/Gradio.")
    bullet(doc, "LLM-assisted natural language explanation (optional).")

    doc.add_page_break()

    # =========================================================
    # PHẦN 13: TÓM TẮT NHANH
    # =========================================================
    h1(doc, "PHẦN 13. BẢNG TÓM TẮT NHANH")

    add_multi_col_table(
        doc,
        rows=[
            ("Đề tài", "A Multimodal Deep Learning Approach for Short-Term Blood Glucose Forecasting in Type 1 Diabetes"),
            ("Bài toán", "Cho lịch sử 120 phút (24 bước × 17 features) + 16 static patient features, dự báo glucose tại 30/60/90 phút tới."),
            ("Loại bài toán ML", "Supervised multi-task regression, multimodal, multi-horizon."),
            ("Dataset", "HUPA-UCM (Hidalgo 2024). 25 bệnh nhân T1D, 309,392 timestep 5-phút, 4 năm 2018-2022."),
            ("Modalities", "Glucose (mg/dL), calories, heart_rate, steps (Fitbit), basal_rate, bolus, carb_input (pump/MDI)."),
            ("Static features", "16 cột (sau pruning 2 vòng từ 25): 4 clinical (HbA1c, age, BMI, dx_time) + 7 derived (train-only) + 3 modality flags + 2 one-hot."),
            ("Dynamic features", "17 cột (sau pruning 2 vòng từ 34): 6 glucose family + 2 raw (HR, basal) + 3 rolling đại diện + 2 pharm (IOB/COB) + 2 time cyclical + 1 censoring + 1 coverage."),
            ("Horizons", "30, 60, 90 phút (t+6, t+12, t+18 bin)."),
            ("Split", "Per-patient chronological 70/15/15, buffer 18-row tại boundary chống leakage."),
            ("Train cap", "Max 5000 sequences/patient trên TRAIN (adaptive stride). Val/test stride=1."),
            ("Sequences", "Train 68,395 / Val 45,382 / Test 45,395 = 159,172 total."),
            ("Scaling", "Continuous → per-subject Z-score (fit train only). Sparse events → log1p + global Z-score."),
            ("Trạng thái hiện tại", "Step 0-4 hoàn thành. Pipeline produce sequence bundle ready cho model. Step 5 (baselines) là kế tiếp."),
            ("Đóng góp dự kiến", "(1) Multimodal nghiêm túc, (2) Patient embedding, (3) Cross-attention fusion, (4) Methodology cho missing modality + sensor caps, (5) XAI + uncertainty + app demo."),
        ],
        headers=["Khía cạnh", "Tóm tắt"],
        widths_cm=[4.5, 11.5],
    )

    para(doc, "")
    info_box(
        doc,
        "Báo cáo này tóm tắt tiến độ tại 18/05/2026. Sau Step 4, pipeline đã sẵn sàng cho việc training baseline "
        "ladder (Step 5). Mọi artefact đều regeneratable bằng cách chạy lại các script trong src/ với SEED=42 cố định.",
    )

    # Save with fallback if the file is open in Word.
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUTPUT))
        target = OUTPUT
    except PermissionError:
        import time as _t
        stamp = _t.strftime("%Y%m%d_%H%M%S")
        target = OUTPUT.with_name(f"{OUTPUT.stem}.{stamp}{OUTPUT.suffix}")
        doc.save(str(target))
        print(f"[WARN] {OUTPUT.name} locked — wrote {target.name} instead.")
    print(f"Saved: {target}")
    print(f"Size: {target.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
