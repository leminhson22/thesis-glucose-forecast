"""Generate the Vietnamese feature-selection Word report (Step 4 §6).

This is the companion to ``feature_catalogue.docx``: it presents the formal
SKILL.md §4.5 feature-selection analysis as an advisor-facing document.

Contents
--------
1. Mục tiêu và phương pháp.
2. 3 signal đánh giá (Spearman, MI, RF permutation) — định nghĩa.
3. Bảng top 20 dynamic features (toàn bộ 3 metric + composite rank).
4. Bảng full 25 static features.
5. Bottom 10 features (candidates ablation).
6. Performance của 3 features mới ở Step 4 (glucose_60m_std, IOB, COB).
7. Quyết định lựa chọn cuối cùng.
8. Hình ảnh embedded (2 figures).

Run from project root:
    python src/generate_feature_selection_report.py
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
SELECTION_CSV = PROJECT_ROOT / "outputs" / "tables" / "hupa_feature_selection.csv"
FIG_DYN = PROJECT_ROOT / "outputs" / "figures" / "04_feature_selection_dynamic.png"
FIG_STAT = PROJECT_ROOT / "outputs" / "figures" / "04_feature_selection_static.png"
OUTPUT = PROJECT_ROOT / "reports" / "feature_selection_report.docx"


# ---------------------------------------------------------------------------
# Docx helpers (duplicated from generate_feature_report.py for self-containment)
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_ranking_table(doc: Document, rows: list[list[str]], headers: list[str], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    for col, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[col].width = Cm(w)

    for col, h in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row_data in enumerate(rows, start=1):
        for col, value in enumerate(row_data):
            cell = table.rows[i].cells[col]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def fmt(x: float, digits: int = 3) -> str:
    if pd.isna(x):
        return "—"
    return f"{x:+.{digits}f}" if abs(x) > 0 else f"0.{'0' * digits}"


def fmt_abs(x: float, digits: int = 3) -> str:
    return f"{x:.{digits}f}" if not pd.isna(x) else "—"


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------


def build() -> None:
    df = pd.read_csv(SELECTION_CSV)
    df["composite_rank"] = df["rank_composite"].astype(float)
    df_dyn = df[df["kind"] == "dynamic"].sort_values("composite_rank").reset_index(drop=True)
    df_stat = df[df["kind"] == "static"].sort_values("composite_rank").reset_index(drop=True)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ===== Title =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO FEATURE SELECTION")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Step 4 — SKILL.md §4.5 — Ranking 34 dynamic + 25 static features bằng 3 signal orthogonal")
    run.italic = True
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Luận văn tốt nghiệp — HUPA-UCM glucose forecasting\n"
        "Target: target_60m (glucose ở t+60 phút)\n"
        "Sample size: 20,000 anchors từ train portion (sub-sampled deterministic)"
    )
    run.font.size = Pt(10)
    doc.add_paragraph()

    # ===== 1. Mục tiêu =====
    add_heading(doc, "1. Mục tiêu của feature selection", level=1)
    add_para(
        doc,
        "Sau khi pipeline tạo 34 dynamic + 25 static = 59 feature, ta cần đánh giá độ informative "
        "của từng feature để (a) xác nhận thiết kế feature là hợp lý, (b) phát hiện feature dư thừa "
        "có thể drop/ablate, (c) cung cấp bằng chứng định lượng cho lựa chọn kiến trúc model ở §7.",
    )
    add_para(
        doc,
        "SKILL.md §4.5 yêu cầu 3 signal: Spearman correlation, mutual information, và "
        "permutation importance từ tree baseline. Đây là 3 góc nhìn khác nhau — mỗi cái bắt được "
        "khía cạnh khác nhau của 'tính quan trọng', nên kết hợp 3 cái cho ranking bền vững hơn 1 cái duy nhất.",
    )

    # ===== 2. Methodology =====
    add_heading(doc, "2. Phương pháp 3 signals", level=1)

    add_heading(doc, "2.1 Spearman rank correlation", level=2)
    add_para(
        doc,
        "Đo độ kết hợp monotonic giữa feature và target_60m. Không yêu cầu quan hệ tuyến tính — "
        "chỉ cần khi feature tăng thì target có xu hướng tăng (hoặc giảm) đều đặn.",
    )
    add_para(
        doc,
        "Công thức: spearman_r = Pearson của (rank của X) với (rank của Y). Giá trị ∈ [-1, +1].",
    )
    add_para(
        doc,
        "Ưu điểm: robust với outlier, không nhạy với scale. Nhược điểm: chỉ bắt được monotonic. "
        "Ví dụ feature 'hour_sin' (sinusoid theo giờ) sẽ có Spearman thấp vì quan hệ là cyclical, không monotonic.",
    )

    add_heading(doc, "2.2 Mutual Information (MI)", level=2)
    add_para(
        doc,
        "Đo lượng thông tin chung giữa feature X và target Y theo entropy: "
        "MI(X;Y) = H(Y) - H(Y|X). Bắt được CỌI LOẠI quan hệ — tuyến tính, non-linear, non-monotonic, "
        "kể cả categorical-vs-continuous.",
    )
    add_para(
        doc,
        "Implementation: sklearn.feature_selection.mutual_info_regression — k-NN entropy estimator của Kraskov 2004.",
    )
    add_para(
        doc,
        "Ưu điểm: capture được mọi pattern, kể cả non-monotonic mà Spearman không thấy. "
        "Nhược điểm: estimate có variance cao trên sample nhỏ; không có scale chuẩn (không phải [-1,1]).",
    )

    add_heading(doc, "2.3 Random Forest Permutation Importance", level=2)
    add_para(
        doc,
        "Train một Random Forest baseline (n_estimators=80, max_depth=12) trên train set, "
        "sau đó: với mỗi feature, shuffle ngẫu nhiên cột đó (phá huỷ relationship với target) "
        "và đo mức giảm R² của model. Drop càng lớn → feature càng quan trọng.",
    )
    add_para(
        doc,
        "Khác biệt cốt lõi: capture được cross-feature interaction. "
        "Ví dụ subject_mean_glucose × glucose hiện tại có thể không informative khi xét riêng "
        "nhưng có vai trò 'normalization' khi kết hợp với feature khác — chỉ permutation importance bắt được điều này.",
    )
    add_para(
        doc,
        "RF baseline đạt R² train = 0.835 trên 20,000 sample — đủ mạnh để các permutation importance scores có ý nghĩa.",
    )

    add_heading(doc, "2.4 Composite ranking", level=2)
    add_para(
        doc,
        "Mỗi feature được xếp hạng theo 3 signal (rank 1 = mạnh nhất). "
        "Composite rank = trung bình 3 rank (unweighted). Cho phép xử lý fair: feature mạnh nhất ở 1 signal "
        "nhưng yếu ở 2 signal khác sẽ ranked thấp hơn feature 'cân đối' top 5 ở cả 3 signal.",
    )

    # ===== 3. Top 20 dynamic =====
    doc.add_page_break()
    add_heading(doc, "3. Top 20 Dynamic Features (sắp theo composite rank)", level=1)
    add_para(
        doc,
        "Dynamic features được đánh giá tại anchor timestep (bước cuối lookback). "
        "Đây là 'giá trị hiện tại' của feature tại thời điểm dự báo — instance informative nhất của lookback.",
    )

    headers = ["Hạng", "Tên feature", "Spearman r", "|Spearman|", "MI", "RF perm imp", "Composite"]
    widths = [1.2, 4.5, 2.0, 1.8, 1.8, 2.4, 2.0]
    rows = []
    for i, r in df_dyn.head(20).iterrows():
        rows.append([
            f"{i+1}",
            r["feature_short"],
            fmt(r["spearman_r"]),
            fmt_abs(r["spearman_abs"]),
            fmt_abs(r["mutual_information"]),
            fmt_abs(r["permutation_importance"], 4),
            f"{r['composite_rank']:.1f}",
        ])
    add_ranking_table(doc, rows, headers, widths)
    doc.add_paragraph()

    add_para(doc, "Nhận xét nhanh về top dynamic:", bold=True)
    add_para(
        doc,
        "• `glucose` raw (rank 1) hoàn toàn dominate — Spearman +0.69 gấp đôi feature mạnh thứ 2, "
        "permutation importance 1.07 trong khi feature thứ 2 chỉ 0.19. Điều này xác nhận EDA §4.4 "
        "rằng glucose hiện tại là tín hiệu autoregressive trọng yếu nhất.",
    )
    add_para(
        doc,
        "• 4 trong top 7 là rolling means của glucose ở các span 30/60/120 phút — xác nhận multi-scale "
        "smoothing là kiến trúc đúng.",
    )
    add_para(
        doc,
        "• `glucose_60m_std` (rank 16) và `glucose_velocity` (rank 9) cho thấy 2nd-order info "
        "(volatility + first derivative) đều có giá trị độc lập với mean.",
    )

    # ===== 4. Static =====
    doc.add_page_break()
    add_heading(doc, "4. Toàn bộ 25 Static Features (sắp theo composite rank)", level=1)
    add_para(
        doc,
        "Static features được broadcast vào mỗi anchor — mọi sample của 1 patient có cùng giá trị static. "
        "Selection trên static đo: 'trong 25 patient, feature này có giúp phân biệt mức glucose tương lai không?'",
    )

    rows = []
    for i, r in df_stat.iterrows():
        rows.append([
            f"{i+1}",
            r["feature_short"],
            fmt(r["spearman_r"]),
            fmt_abs(r["spearman_abs"]),
            fmt_abs(r["mutual_information"]),
            fmt_abs(r["permutation_importance"], 4),
            f"{r['composite_rank']:.1f}",
        ])
    add_ranking_table(doc, rows, headers, widths)
    doc.add_paragraph()

    add_para(doc, "Nhận xét quan trọng:", bold=True)
    add_para(
        doc,
        "• 5/10 top composite features (xét cả dynamic và static) là STATIC: subject_mean_glucose (rank 2 overall), "
        "subject_std_glucose, subject_hyper_pct, subject_tir_pct, subject_hypo_pct. "
        "Đây là bằng chứng định lượng MẠNH NHẤT cho việc dùng patient-embedding branch — không phải sự lựa chọn stylistic.",
    )
    add_para(
        doc,
        "• Các derived static (TIR%, hypo%, std...) tính từ train portion ONLY — không có leakage. "
        "Việc chúng ranking cao xác nhận: pattern lịch sử của patient TRONG TRAIN dự báo tốt cho FUTURE của patient.",
    )
    add_para(
        doc,
        "• Các metadata lâm sàng (HbA1c, bmi, age, dx_time) ranking trung bình — informative nhưng "
        "không bằng behavioural fingerprint. Điều này cho thấy 'hành vi quan sát được' mạnh hơn 'đặc điểm bệnh lý'.",
    )

    # ===== 5. Bottom features =====
    doc.add_page_break()
    add_heading(doc, "5. Bottom 10 Features (ablation candidates)", level=1)
    add_para(
        doc,
        "Bottom features không tự động drop — nhưng được FLAG là candidates cho ablation study ở §8 report.",
    )

    bottom = df.sort_values("composite_rank", ascending=False).head(10)
    rows = []
    for _, r in bottom.iterrows():
        rows.append([
            r["kind"],
            r["feature_short"],
            fmt(r["spearman_r"]),
            fmt_abs(r["spearman_abs"]),
            fmt_abs(r["mutual_information"]),
            fmt_abs(r["permutation_importance"], 4),
            f"{r['composite_rank']:.1f}",
        ])
    headers_b = ["Kind", "Tên feature", "Spearman r", "|Spearman|", "MI", "RF perm imp", "Composite"]
    widths_b = [1.5, 4.2, 2.0, 1.8, 1.8, 2.4, 2.0]
    add_ranking_table(doc, rows, headers_b, widths_b)
    doc.add_paragraph()

    add_para(doc, "Phân tích 3 nhóm bottom:", bold=True)
    add_para(
        doc,
        "Nhóm 1 — Dynamic modality flags (basal/bolus/carb_available): "
        "Cờ này constant per-patient (mọi row của 1 patient cùng giá trị), nên ở per-anchor level "
        "không có variance để bắt signal. VAI TRÒ thực của chúng là ở patient-level — đã có trong static version. "
        "Đây là REDUNDANCY THỰC SỰ giữa dynamic và static.",
    )
    add_para(
        doc,
        "Nhóm 2 — Raw event streams (carb_input, bolus_volume_delivered, steps raw): "
        "Distribution dominated by zeros (>97% bin = 0). Marginal info kém nhưng giữ TIMING PRECISION — "
        "model có thể biết CHÍNH XÁC bolus tiêm ở bin nào, không phải chỉ rolling sum smear. "
        "Khuyến nghị: giữ, ablate ở §8 để verify timing có giá trị riêng.",
    )
    add_para(
        doc,
        "Nhóm 3 — One-hot redundancy (treatment_MDI vs CSII, gender_Male vs Female): "
        "Mỗi cặp là one-hot 2 cột tổng = 1. Drop một cột không làm mất info nhưng làm asymmetric. "
        "Khuyến nghị: giữ cặp, model dense layer không bị ảnh hưởng bởi colinearity.",
    )

    # ===== 6. New features Step 4 =====
    add_heading(doc, "6. Performance của 3 features mới ở Step 4", level=1)
    add_para(
        doc,
        "3 features được thêm ở Step 4 (sau khi Step 3 đã có 31 dynamic + 25 static). "
        "Kiểm chứng xem có justify được không.",
    )

    new_feats = ["glucose_60m_std", "insulin_on_board", "carbs_on_board"]
    rows = []
    for feat in new_feats:
        r = df[df["feature_short"] == feat].iloc[0]
        rank_overall = df.sort_values("composite_rank").reset_index(drop=True).index[
            df.sort_values("composite_rank").reset_index(drop=True)["feature_short"] == feat
        ].tolist()[0] + 1
        rows.append([
            feat,
            f"{rank_overall} / {len(df)}",
            fmt(r["spearman_r"]),
            fmt_abs(r["mutual_information"]),
            fmt_abs(r["permutation_importance"], 4),
            f"{r['composite_rank']:.1f}",
        ])
    headers_new = ["Feature", "Rank overall", "Spearman r", "MI", "RF perm imp", "Composite"]
    widths_new = [3.5, 2.3, 2.2, 2.0, 2.6, 2.0]
    add_ranking_table(doc, rows, headers_new, widths_new)
    doc.add_paragraph()

    add_para(
        doc,
        "Kết luận: cả 3 đều rank middle pack (16-41 trong 59), KHÔNG phải feature yếu, "
        "và cả 3 đều có MI / perm imp > 0 → có đóng góp predictive thực sự. "
        "Đặc biệt glucose_60m_std rank 16 là mạnh hơn nhiều rolling features của events.",
    )

    # ===== 7. Final decisions =====
    add_heading(doc, "7. Quyết định selection cuối cùng", level=1)
    add_para(doc, "DROP (cho final modelling): 0 features.", bold=True)
    add_para(
        doc,
        "Lý do giữ all 34 + 25: chi phí marginal nhỏ (3 cờ dynamic chỉ thêm 72 số / sample = 9% overhead), "
        "trong khi việc DROP dựa trên một analysis duy nhất là quyết định kém khoa học hơn việc ABLATE ở §8.",
    )
    add_para(doc, "FLAG (cho ablation §8):", bold=True)
    add_para(
        doc,
        "• Dynamic basal/bolus/carb_available: ablate với 'all 34' vs '31 dynamic + 3 static availability only'.",
    )
    add_para(
        doc,
        "• Raw event streams (carb_input, bolus_volume_delivered, steps): ablate '34 with raw' vs '31 with rolling only' "
        "để verify timing precision đáng giữ.",
    )
    add_para(doc, "KEEP firmly (không ablate):", bold=True)
    add_para(
        doc,
        "• Top 10 features (rank composite ≤ 10.7): không có lý do gì để remove.",
    )
    add_para(
        doc,
        "• 3 features mới Step 4 (IOB, COB, glucose_60m_std): đều middle-pack, có literature backing.",
    )

    # ===== 8. Embedded figures =====
    doc.add_page_break()
    add_heading(doc, "8. Hình ảnh trực quan", level=1)
    add_para(doc, "Top 20 dynamic features — 3 signal song song:")
    if FIG_DYN.exists():
        doc.add_picture(str(FIG_DYN), width=Cm(16))
    add_para(doc, "Toàn bộ 25 static features — 3 signal song song:")
    if FIG_STAT.exists():
        doc.add_picture(str(FIG_STAT), width=Cm(16))

    add_heading(doc, "9. Kết luận và bước tiếp theo", level=1)
    add_para(
        doc,
        "Feature selection xác nhận thiết kế Step 3 + Step 4 là hợp lý: không có feature nào nên drop "
        "ngay, nhưng đã định lượng được redundancy của 3 dynamic availability flags và pattern timing "
        "vs aggregate của event streams. Hai pattern này được flag là ablation candidates trong §8.",
    )
    add_para(
        doc,
        "Bằng chứng quan trọng nhất cho §7 (Modelling Strategy): "
        "static features chiếm 5/10 top composite ranks, → kiến trúc 2-branch với patient-embedding "
        "không phải lựa chọn stylistic mà là empirically justified.",
    )
    add_para(
        doc,
        "Pipeline sẵn sàng cho Step 5 (baseline ladder): persistence → Ridge → RF/XGBoost → LSTM/GRU.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
