"""Generate the Vietnamese feature-catalogue Word document (38 features).

This is the post-pruning version. After the Step 4-revisit dropped 21
redundant columns, the final modelling configuration is 19 dynamic +
19 static features. Each entry has the same fields:
    - Tên cột
    - Ý nghĩa vật lý / lâm sàng
    - Công thức / cách tính
    - Mục đích sử dụng cho model
    - Bằng chứng / rationale
    - Nguồn gốc (kế thừa từ pipeline tác giả / thesis tạo mới)

Run from project root:
    python src/generate_feature_report.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "reports" / "feature_catalogue.docx"


# ---------------------------------------------------------------------------
# Feature definitions (FINAL 38 features after Step 4-revisit pruning)
# ---------------------------------------------------------------------------

DYNAMIC_GROUPS: list[tuple[str, str, list[tuple[str, str, str, str, str, str, str]]]] = [
    (
        "Nhóm A — Kênh glucose và các feature dẫn xuất (6 cột)",
        "Nhóm trọng yếu nhất — biến cần dự báo chính là glucose tại horizon 30/60/90 phút. "
        "Clinical-lens revisit đã drop glucose_acceleration vì không CGM device nào show 2nd derivative "
        "và không guideline lâm sàng nào dùng. Velocity đã capture actionable trend.",
        [
            (
                "glucose",
                "float32 (mg/dL, đã Z-score per-subject)",
                "Nồng độ glucose mô kẽ tại bin 5-min, đo bởi FreeStyle Libre 2.",
                "Kế thừa từ glUCModel: round 5-min → subsample 15-min → linear interp về 5-min (max 1h gap). "
                "Áp Z-score per-subject tại Step 3.",
                "Tín hiệu autoregressive trọng yếu nhất. Spearman r=+0.693 với target_60m. "
                "Permutation importance gấp đôi feature mạnh thứ 2.",
                "ACF half-life 130 phút (EDA §4.4). Feature selection rank 1/38.",
                "Kế thừa.",
            ),
            (
                "glucose_30m_mean",
                "float32 (mg/dL, đã Z-score)",
                "Trung bình glucose trong 30 phút gần nhất (6 bin).",
                "rolling(window=6, min_periods=1).mean() per-participant.",
                "Khử nhiễu sensor cao tần, cung cấp 'mức nền' ngắn hạn.",
                "EDA §4.7: SD ngắn hạn glucose có thể đạt 9 mg/dL trong 5 phút khi đường dốc. Rank 3/38.",
                "Mới.",
            ),
            (
                "glucose_60m_mean",
                "float32 (mg/dL, đã Z-score)",
                "Trung bình glucose trong 60 phút gần nhất (12 bin).",
                "rolling(window=12, min_periods=1).mean().",
                "Trend trung hạn. Cùng 30m_mean tạo cặp short/medium giúp model so sánh nhịp thay đổi.",
                "Horizon 60 phút trùng target_60m. Rank 7/38.",
                "Mới.",
            ),
            (
                "glucose_120m_mean",
                "float32 (mg/dL, đã Z-score)",
                "Trung bình glucose trong 120 phút gần nhất (24 bin = full lookback).",
                "rolling(window=24, min_periods=1).mean().",
                "'Baseline' của patient trong window. Phân biệt hypo-rebound vs down-drift.",
                "ACF lag 120 phút trung bình 0.5 (EDA §4.4). Rank 4/38.",
                "Mới.",
            ),
            (
                "glucose_60m_std",
                "float32 (mg/dL, đã Z-score)",
                "Độ lệch chuẩn glucose 60 phút gần nhất — proxy volatility ngắn hạn.",
                "rolling(window=12, min_periods=2).std(). Row đầu fillna(0).",
                "Phân biệt 'ổn định' vs 'dao động mạnh'. NN khó tự tính std từ raw lookback. "
                "Hữu ích cho uncertainty quantification ở §10.",
                "Battelino 2019 international consensus: CGV (CGM variability) là metric lâm sàng chuẩn.",
                "Mới (Step 4).",
            ),
            (
                "glucose_velocity",
                "float32 (mg/dL/phút, đã Z-score)",
                "Tốc độ thay đổi glucose — first derivative.",
                "diff(glucose) / 5. Row đầu fillna(0).",
                "Velocity âm → cảnh báo hypo 30 phút tới. Velocity dương sau bữa → cảnh báo hyper. "
                "Trên CGM display, đây là 'trend arrows' (↑↑, ↑, →, ↓, ↓↓) — clinicians sử dụng trực tiếp.",
                "EDA §4.7: p95 |velocity| = 9 mg/dL/phút. Rank 7/33 (top 21%).",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm B — Kênh raw từ thiết bị (2 cột)",
        "Đã giảm từ 6 → 2 sau Step 4-revisit. Đã drop: calories (derived từ HR+steps), steps (raw mostly zero, "
        "rolling_150m thay thế), bolus_volume_delivered (recoverable từ IOB), carb_input (recoverable từ COB). "
        "Còn lại heart_rate và basal_rate vì cả hai đều có timing precision riêng.",
        [
            (
                "heart_rate",
                "float32 (bpm, đã Z-score per-subject)",
                "Nhịp tim tại bin 5-min, đo bởi Fitbit Ionic.",
                "Kế thừa: round 5-min + linear interpolation. Áp Z-score per-subject (baseline HR khác biệt rõ giữa patients).",
                "HR cao → stress/exercise → epinephrine → glucose tăng. HR thấp ngủ → glucose ổn. "
                "Sudden HR spike (raw value) bắt được mà rolling mean smooth ra.",
                "Bequette 2017, Acciaroli 2018. Botella-Serrano 2023 dùng HUPA cho sleep×glycemic study.",
                "Kế thừa.",
            ),
            (
                "basal_rate",
                "float32 (U/5-min, đã log1p + global Z-score)",
                "Liều insulin nền truyền/tiêm trong bin 5-min.",
                "CSII: sum rate. MDI: liều long-acting chia 288. Áp log1p + global Z-score.",
                "Insulin nền hạ glucose chậm 6-24h. Model cần để hiểu glucose giảm dần khi không có bolus.",
                "Cặp với basal_coverage_24h + basal_available (static) để phân biệt 'không tiêm' vs 'không ghi nhận'.",
                "Kế thừa.",
            ),
        ],
    ),
    (
        "Nhóm C — Rolling aggregates đại diện (3 cột)",
        "Đã giảm từ 9 → 3. Drop bolus_30m, bolus_180m, carb_60m, carb_180m (IOB/COB cover), steps_30m (subset), "
        "calories_30m_sum (calories derived từ Fitbit HR+steps, không clinical cho glucose management).",
        [
            (
                "bolus_60m_sum",
                "float32 (U, đã log1p + global Z-score)",
                "Tổng bolus insulin tiêm trong 60 phút gần nhất (12 bin) — span đại diện.",
                "rolling(window=12, min_periods=1).sum().",
                "Span 60m là pha peak của rapid-acting insulin. Giữ explicit cumulative cho tree baselines "
                "không thể tự reconstruct IIR filter từ lookback.",
                "Mathieu 2017 pharmacokinetics. Span 30m và 180m đã được drop, IOB cover đầy đủ.",
                "Mới.",
            ),
            (
                "steps_150m_sum",
                "float32 (số bước, đã log1p + global Z-score)",
                "Tổng bước trong 150 phút gần nhất (30 bin).",
                "rolling(window=30, min_periods=1).sum().",
                "Bắt session vận động dài + 'after-exercise effect' kéo dài 1-2h sau exercise "
                "(glucose tiếp tục giảm do bù glycogen).",
                "Riddell 2017: post-exercise insulin sensitivity tăng kéo dài. Span 30m drop vì subset.",
                "Mới.",
            ),
            (
                "heart_rate_30m_mean",
                "float32 (bpm, đã Z-score per-subject)",
                "HR trung bình 30 phút — smoothed baseline.",
                "rolling(window=6, min_periods=1).mean().",
                "Phân biệt 'stress kéo dài' (mean cao) vs 'spike thoáng qua' (raw cao nhưng mean bình thường). "
                "Sleep state: mean liên tục dưới baseline → glucose có xu hướng giảm.",
                "Cặp với raw HR cho timing precision của bursts.",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm C2 — Pharmacokinetic-decay features (2 cột)",
        "Đại diện 'on board' cho insulin và carb. Đã thay thế cho 5 rolling sums (bolus 3 span + carb 2 span) "
        "vì IOB/COB là 1-pole IIR filter — biết IOB[t-23:t] có thể tính ngược raw bolus exactly. Khác rolling sum "
        "(rectangular kernel) ở chỗ kernel exponential decay matching dược động học thực.",
        [
            (
                "insulin_on_board",
                "float32 (U, đã log1p + global Z-score)",
                "Insulin On Board — lượng bolus VẪN CÒN tác dụng tại thời điểm hiện tại.",
                "IOB[t] = α·IOB[t-1] + bolus[t], với α = exp(-5/75) = 0.9355. Tính per-patient, reset =0 ở row đầu.",
                "Kernel pharmacokinetic mà tree-based model không thể tự suy ra. Đặc biệt giúp dự báo "
                "hypo muộn (60-90 phút sau bolus).",
                "τ=75 phút match action time rapid-acting analogs (Mathieu 2017). Sau 5τ=375 phút residual <1%.",
                "Mới (Step 4).",
            ),
            (
                "carbs_on_board",
                "float32 (servings, đã log1p + global Z-score)",
                "Carbs On Board — lượng carb VẪN ĐANG hấp thu.",
                "COB[t] = α·COB[t-1] + carb[t], với α = exp(-5/60) = 0.9200.",
                "Cùng IOB tạo cặp đối lập (insulin hạ vs carb tăng). Model học net effect bằng phép trừ.",
                "τ=60 phút theo Hovorka 2004 first-order absorption.",
                "Mới (Step 4).",
            ),
        ],
    ),
    (
        "Nhóm D — Encoding chu kỳ thời gian (2 cột)",
        "Đã giảm từ 4 → 2. Drop dayofweek_sin/cos vì weekend effect chỉ 3-4 điểm % TIR (yếu). Giữ hour vì "
        "circadian effect mạnh (peak 07-10h, nadir 03-05h).",
        [
            (
                "hour_sin",
                "float32 ([-1, 1], pass-through)",
                "Sin component giờ trong ngày, chu kỳ 24h.",
                "sin(2π × hour / 24).",
                "Cặp với hour_cos tạo embedding 2-D liên tục của giờ. Tránh discontinuity tại 23h59→00h00.",
                "EDA §4.5: glucose peak 07-10h ~165 mg/dL, nadir 03-05h ~135 mg/dL.",
                "Mới.",
            ),
            (
                "hour_cos",
                "float32 ([-1, 1], pass-through)",
                "Cos component giờ trong ngày.",
                "cos(2π × hour / 24).",
                "Cặp đôi với hour_sin.",
                "Same as above.",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm E — Cờ kiểm duyệt sensor (1 cột)",
        "Đã giảm từ 2 → 1. Drop glucose_high_extreme (chỉ 0.04% rows, không đủ frequency). Giữ glucose_low_cap "
        "vì hypoglycemia clinically critical nhất.",
        [
            (
                "glucose_low_cap",
                "int8 ({0, 1}, pass-through)",
                "Cờ nhị phân: 1 nếu glucose ≤ 40 (FreeStyle Libre 2 LO).",
                "(glucose <= 40).astype(int8).",
                "Cho loss có Tobit-style mask: tin tưởng kém ở censored values. XAI analysis behavior xung quanh LO.",
                "EDA §3.6.4: 0.38% bin globally, peak 5.82% (HUPA0002P).",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm F — Coverage rolling (1 cột)",
        "Đã giảm từ 4 → 1. Drop 3 dynamic availability flags (basal/bolus/carb_available) vì constant per-patient, "
        "MI=0, đã duplicate trong static branch. Giữ basal_coverage_24h vì time-varying.",
        [
            (
                "basal_coverage_24h",
                "float32 ([0, 1], pass-through)",
                "Tỷ lệ bin có basal>0 trong 24h gần nhất.",
                "rolling(288, min_periods=1).mean() của (basal_rate > 0).",
                "Phân biệt 3 chế độ: CSII liên tục (~1.0), MDI sparse (~0.05), partial coverage (0.4-0.66). "
                "Khác static flag basal_available ở chỗ time-varying.",
                "EDA §3.6.3: 4 patient (HUPA0024/26/27/28) có basal coverage 40-66%.",
                "Mới.",
            ),
        ],
    ),
]


STATIC_GROUPS: list[tuple[str, str, list[tuple[str, str, str, str, str, str, str]]]] = [
    (
        "Nhóm I — Metadata lâm sàng (4 cột)",
        "Đã giảm từ 6 → 4. Drop weight_kg, height_cm vì BMI = weight/(height/100)² là deterministic function. "
        "BMI rank cao hơn cả 2 cộng lại trong selection.",
        [
            (
                "hba1c_pct",
                "float32 (%, đã Z-score)",
                "HbA1c — đường huyết trung bình 3 tháng. Range gốc 6.0-9.7%.",
                "Đọc từ xlsx, Z-score across 25 patient.",
                "Chỉ số kiểm soát glucose nền. HbA1c cao → distribution lệch hyperglycemia.",
                "ADA Standards of Care. Tena 2021, Parra 2024 đều dùng.",
                "Kế thừa.",
            ),
            (
                "age_years",
                "float32 (năm, đã Z-score)",
                "Tuổi. Range 18.0-61.8.",
                "Đọc từ xlsx, Z-score.",
                "Patient trẻ insulin sensitivity cao hơn. Tuổi correlate với complication history.",
                "Standard T1D covariate.",
                "Kế thừa.",
            ),
            (
                "dx_time_years",
                "float32 (năm, đã Z-score)",
                "Thời gian mắc T1DM. Range 0.8-39.5.",
                "Đọc từ xlsx, Z-score.",
                "Lâu năm → không còn C-peptide → glucose volatile. Mới mắc → 'honeymoon period'.",
                "DCCT/EDIC: complications risk ~1.5%/year.",
                "Kế thừa.",
            ),
            (
                "bmi",
                "float32 (kg/m², đã Z-score)",
                "Body Mass Index.",
                "weight_kg / (height_cm/100)². Z-score across 25 patient.",
                "Insulin resistance correlate với BMI. 8/25 patient BMI > 25 (overweight).",
                "Hovorka 2004 dùng BMI làm patient covariate. Rank 15/38 (top 40%).",
                "Mới (tự tính từ raw, drop weight+height sau).",
            ),
        ],
    ),
    (
        "Nhóm II — Behavioral fingerprint từ TRAIN PORTION ONLY (7 cột)",
        "Đã giảm từ 12 → 7. Drop: subject_tir_pct (math redundant), mean_daily_steps (correlate steps_active_pct), "
        "carb_events_per_day (compliance proxy không phải biology), data_duration_days (data artifact), "
        "basal_recording_pct (correlate ~0.9 với treatment_CSII, basal_available encode signal). Tính từ train only.",
        [
            (
                "subject_mean_glucose",
                "float32 (đã Z-score)",
                "Glucose trung bình của patient từ train rows.",
                "df_train.groupby('participant_id')['glucose'].mean(). Z-score.",
                "Baseline glucose của patient. Cao → cần dự báo conservative hơn.",
                "EDA §4.7: range 113-201 mg/dL. **Rank 2/38** (chỉ sau glucose raw).",
                "Mới.",
            ),
            (
                "subject_std_glucose",
                "float32 (đã Z-score)",
                "Glucose variability của patient.",
                "Tương tự nhưng .std(). Z-score.",
                "Patient có std cao là khó dự báo hơn — model output uncertainty rộng hơn.",
                "Battelino 2019. Rank 6/38.",
                "Mới.",
            ),
            (
                "subject_hypo_pct",
                "float32 (đã Z-score)",
                "% thời gian hypo (<70 mg/dL) trong train portion.",
                "(glucose < 70).mean() × 100. Z-score.",
                "Patient nhiều hypo → model dự báo thận trọng zone thấp.",
                "Rank 9/38.",
                "Mới.",
            ),
            (
                "subject_hyper_pct",
                "float32 (đã Z-score)",
                "% thời gian hyper (>180) trong train portion.",
                "(glucose > 180).mean() × 100. Z-score.",
                "Bổ trợ hypo_pct. Cùng nhau xác định TIR (đã drop tir_pct redundant).",
                "Rank 5/38.",
                "Mới.",
            ),
            (
                "bolus_events_per_day",
                "float32 (đã Z-score)",
                "Số lần tiêm bolus trung bình mỗi ngày.",
                "(bolus > 0).sum() / data_duration_days. Z-score.",
                "Lifestyle compliance proxy. Pump users 4-7 bolus/ngày; MDI 3-4.",
                "Rank 11/38.",
                "Mới.",
            ),
            (
                "steps_active_pct",
                "float32 (đã Z-score)",
                "% bin có steps > 99 ('đi bộ tích cực').",
                "(steps > 99).mean() × 100. Z-score.",
                "Activity intensity profile. Active patients có insulin sensitivity tốt hơn.",
                "Threshold 99 steps/5min ≈ 20 steps/min ≈ walking. Rank 10/38.",
                "Mới.",
            ),
            (
                "mean_heart_rate",
                "float32 (đã Z-score)",
                "HR trung bình của patient.",
                "heart_rate.mean(). Z-score.",
                "Baseline HR khác giữa patients (60-85 bpm). Cần để normalize HR raw.",
                "Standard biometric.",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm III — Modality availability flags (3 cột)",
        "Per-patient binary. Khác dynamic version đã drop — version dynamic constant per-patient nên duplicate. "
        "Giữ static để patient embedding biết toàn cục modality nào có.",
        [
            (
                "basal_available",
                "int8 ({0, 1})",
                "Patient có basal record bất kỳ?",
                "(basal_rate > 0).any() trên toàn timeline patient.",
                "Patient embedding biết modality nào có thể tin. basal_available=0 → mọi giá trị basal_rate=0 là 'không recorded'.",
                "Pitfall #6: 4 patient không có basal.",
                "Mới.",
            ),
            (
                "bolus_available",
                "int8 ({0, 1})",
                "Patient có bolus record bất kỳ?",
                "Tương tự.",
                "Tương tự.",
                "Pitfall #6: 3 patient không có bolus.",
                "Mới.",
            ),
            (
                "carb_available",
                "int8 ({0, 1})",
                "Patient có carb record bất kỳ?",
                "Tương tự.",
                "Tương tự.",
                "Pitfall #6: 3 patient không có carb.",
                "Mới.",
            ),
        ],
    ),
    (
        "Nhóm IV — Demographic one-hot (2 cột)",
        "Đã giảm từ 4 → 2. Drop gender_Male và treatment_MDI vì one-hot pair sum=1, redundant với complement. "
        "Giữ gender_Female và treatment_CSII làm indicator độc lập.",
        [
            (
                "gender_Female",
                "float32 ({0, 1})",
                "Cờ giới tính nữ.",
                "pd.get_dummies(static, columns=['gender'])['gender_Female']. 13/25 patient.",
                "Insulin sensitivity khác theo giới (chu kỳ hormone, mass cơ).",
                "Hovorka 2004 dùng gender. HUPA cohort 52% nữ.",
                "Mới (one-hot, drop Male).",
            ),
            (
                "treatment_CSII",
                "float32 ({0, 1})",
                "Cờ patient dùng pump CSII (continuous subcutaneous insulin infusion).",
                "pd.get_dummies. HUPA0011P override CSII→MDI trước encode (Pitfall #8). 13/25 CSII.",
                "Phân biệt phong cách quản lý insulin: CSII basal liên tục, MDI sparse.",
                "Pickup 2014 T1D research standard covariate.",
                "Mới (one-hot + override).",
            ),
        ],
    ),
]


TARGETS_INFO = [
    (
        "target_30m",
        "float32 (mg/dL, scale gốc)",
        "Glucose tại t+6 bin (= 30 phút sau).",
        "df.groupby('participant_id')['glucose'].shift(-6).",
        "Horizon ngắn — immediate hypo alert. Dễ dự báo nhất.",
        "CLAUDE.md horizon spec.",
        "Mới.",
    ),
    (
        "target_60m",
        "float32 (mg/dL, scale gốc)",
        "Glucose tại t+12 bin (= 60 phút sau).",
        ".shift(-12).",
        "Horizon trung — chuẩn báo cáo CGM forecasting.",
        "CLAUDE.md.",
        "Mới.",
    ),
    (
        "target_90m",
        "float32 (mg/dL, scale gốc)",
        "Glucose tại t+18 bin (= 90 phút sau).",
        ".shift(-18).",
        "Horizon dài — challenging, đo khả năng học long-range dependency.",
        "CLAUDE.md.",
        "Mới.",
    ),
]


# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_feature_table(doc, rows):
    headers = ["Tên", "Kiểu / đơn vị", "Ý nghĩa", "Cách tính", "Mục đích", "Bằng chứng", "Nguồn gốc"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    widths_cm = [2.4, 1.7, 3.0, 3.0, 3.0, 2.6, 1.4]
    for col, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[col].width = Cm(w)
    for col, h in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row_data in enumerate(rows, start=1):
        for col, value in enumerate(row_data):
            cell = table.rows[i].cells[col]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CATALOG FEATURE — PHIÊN BẢN ĐÃ PRUNE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "33 features cuối cùng (17 dynamic + 16 static) sau Step 4 information-theoretic + clinical-lens pruning từ 59 features"
    )
    run.italic = True
    run.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Luận văn tốt nghiệp · HUPA-UCM glucose forecasting · Pipeline sẵn sàng cho Step 5 baseline ladder"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "0. Tổng quan kiến trúc feature sau 2 lần pruning", level=1)
    add_para(
        doc,
        "Pipeline tiền xử lý cuối cùng sinh ra 2 loại input:",
    )
    add_para(
        doc,
        "• Dynamic features (17 cột) — mỗi sample shape (24, 17). Phim 120 phút lookback của patient.",
    )
    add_para(
        doc,
        "• Static features (16 cột) — vector 16-d cố định per patient. CMND của bệnh nhân.",
    )
    add_para(
        doc,
        "Mục đích tài liệu này: thuyết minh chi tiết 33 features được pipeline §5 sinh ra cho model. "
        "Phiên bản gốc (59 features) đã được pruning theo 2 vòng:",
    )
    add_para(
        doc,
        "Vòng 1 (information-theoretic, drop 21): math redundancy (weight/height/Male/MDI), IOB/COB-redundant "
        "(raw bolus/carb + 4 rolling sums), wearable-redundant (raw calories/steps + mean_daily_steps), "
        "statistical (TIR%), weak signal (dayofweek, glucose_high_extreme, 3 dynamic availability flags).",
    )
    add_para(
        doc,
        "Vòng 2 (clinical lens, drop 5): glucose_acceleration (không clinical display dùng 2nd derivative), "
        "calories_30m_sum (derived từ Fitbit HR+steps, không clinical cho glucose), "
        "carb_events_per_day (compliance proxy không phải biology), "
        "data_duration_days (data artifact zero clinical meaning), "
        "basal_recording_pct (correlate ~0.9 với treatment_CSII).",
    )

    add_heading(doc, "Tóm tắt thống kê pipeline output", level=2)
    summary_table = doc.add_table(rows=9, cols=2)
    summary_table.style = "Light Grid Accent 1"
    summary_data = [
        ("Tên thông số", "Giá trị"),
        ("Số patient", "25"),
        ("Lookback window", "24 bước (120 phút)"),
        ("Horizons dự báo", "30, 60, 90 phút"),
        ("Số dynamic features", "17 (giảm từ 34, drop 17 sau 2 vòng pruning)"),
        ("Số static features", "16 (giảm từ 25, drop 9 sau 2 vòng pruning)"),
        ("Tổng features", "33 (giảm 44.1% từ 59 gốc)"),
        ("Tổng sequences", "159,172 (train 68,395 / val 45,382 / test 45,395)"),
        ("Feature budget per sample", "24 × 17 + 16 = 424 (giảm 49.6% từ 841)"),
    ]
    for i, (k, v) in enumerate(summary_data):
        for col, val in enumerate((k, v)):
            cell = summary_table.rows[i].cells[col]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "1F3A5F")

    doc.add_paragraph()

    # PART 1: DYNAMIC
    doc.add_page_break()
    add_heading(doc, "PHẦN 1 — DYNAMIC FEATURES (17 cột × 24 timestep)", level=1)
    add_para(
        doc,
        "Input chuỗi thời gian. Mỗi sample shape (24, 17) trong data/processed/hupa_5min_sequences.npz.",
    )
    add_para(
        doc,
        "Thứ tự cố định trong dynamic_feature_order() (src/preprocessing.py). Notebook downstream lookup index của "
        "feature qua tên trong feature_names_dynamic.",
    )

    for group_name, group_intro, group_rows in DYNAMIC_GROUPS:
        add_heading(doc, group_name, level=2)
        add_para(doc, group_intro)
        add_feature_table(doc, group_rows)
        doc.add_paragraph()

    # PART 2: STATIC
    doc.add_page_break()
    add_heading(doc, "PHẦN 2 — STATIC FEATURES (16 cột × 1 vector per patient)", level=1)
    add_para(
        doc,
        "Vector 16-d per patient, broadcast vào mỗi anchor. Shape (159172, 16) trong npz.",
    )
    add_para(
        doc,
        "Lưu ý quan trọng: 7 cột nhóm II (derived) tính CHỈ TỪ phần TRAIN — tránh leakage. "
        "Patient ở val/test set vẫn có giá trị derived (broadcast từ train portion của chính họ).",
    )

    for group_name, group_intro, group_rows in STATIC_GROUPS:
        add_heading(doc, group_name, level=2)
        add_para(doc, group_intro)
        add_feature_table(doc, group_rows)
        doc.add_paragraph()

    # PART 3: TARGETS
    doc.add_page_break()
    add_heading(doc, "PHẦN 3 — TARGETS Y (3 horizon)", level=1)
    add_para(
        doc,
        "Y shape (159172, 3). Mỗi sequence 3 horizon target ở scale gốc mg/dL (không scaled). "
        "Cho phép tính RMSE/MAE trực tiếp đơn vị lâm sàng.",
    )
    add_feature_table(doc, TARGETS_INFO)
    doc.add_paragraph()

    # PART 4: Closing
    add_heading(doc, "PHẦN 4 — Cách model tiêu thụ input ở Step 5+", level=1)
    add_para(doc, "Baseline tabular (Ridge, RF, XGBoost):", bold=True)
    add_para(
        doc,
        "  X = concat(X_dynamic.reshape(N, 24*17), X_static) → shape (N, 424)\n"
        "  y = y → shape (N, 3) hoặc tách 3 model cho 3 horizon.",
    )
    add_para(doc, "Baseline neural đơn (LSTM/GRU):", bold=True)
    add_para(
        doc,
        "  h_T = LSTM(X_dynamic)[-1] (hidden state cuối)\n"
        "  z = Dense(static_dim)(X_static)\n"
        "  y_hat = Output_head(concat(h_T, z)) → 3 horizon.",
    )
    add_para(doc, "Hybrid CNN-GRU-Attention đề xuất:", bold=True)
    add_para(
        doc,
        "  CNN 1D trên X_dynamic → local pattern (5-15 phút).\n"
        "  GRU bidirectional → long-range dependency.\n"
        "  MLP(X_static) → patient embedding z.\n"
        "  Cross-attention(query=GRU_out, key/value=z) → gate temporal theo patient.\n"
        "  Multi-horizon output head.",
    )

    add_heading(doc, "PHẦN 5 — Sanity check trước Step 5", level=1)
    add_para(doc, "✓ Không NaN trong X_dynamic, X_static, y (verified).")
    add_para(doc, "✓ Không leakage thời gian (test anchor_time > val > train per patient).")
    add_para(doc, "✓ Scalers fit on TRAIN only, persisted ở outputs/models/scalers.json.")
    add_para(doc, "✓ Long-patient capped 5,000 sequences/patient trên train (HUPA0026/0027/0028).")
    add_para(doc, "✓ Pipeline deterministic ở SEED=42, runtime ~50s local.")
    add_para(doc, "✓ Feature selection re-run trên 38 features, top 10 vẫn glucose + 6 static derived.")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
